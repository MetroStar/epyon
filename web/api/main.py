"""
Epyon Web — FastAPI backend
"""
from __future__ import annotations

import asyncio
import io
import json
import logging
import os
import re
import shutil
import time
import zipfile
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, List, Optional

from fastapi import FastAPI, HTTPException, Query, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from . import jobs as job_store
from . import parsers
from . import github_sync
from . import github_metrics
from . import openai_summary
from . import jira_client

# ── Paths ─────────────────────────────────────────────────────
_HERE        = Path(__file__).parent
EPYON_ROOT   = (_HERE / ".." / "..").resolve()
VERSION_FILE = EPYON_ROOT / "VERSION"
EPYON_VERSION = VERSION_FILE.read_text(encoding="utf-8").strip() if VERSION_FILE.exists() else "unknown"
SCRIPTS_DIR  = EPYON_ROOT / "scripts" / "shell"
APPROVED_IMAGES_FILE = EPYON_ROOT / "configuration" / "approved-base-images.conf"
GITHUB_CONFIG_FILE   = _HERE / ".." / "github-config.json"
HIDDEN_APPS_FILE     = EPYON_ROOT / "configuration" / "hidden-apps.json"
JIRA_CONFIG_FILE     = (_HERE / ".." / "data" / "jira-config.json").resolve()
JIRA_TICKETS_FILE    = (_HERE / ".." / "data" / "jira-tickets.json").resolve()
REGISTERED_APPS_FILE = EPYON_ROOT / "configuration" / "registered-apps.json"
MONITORED_APPS_FILE  = EPYON_ROOT / "configuration" / "monitored-apps.json"
STATIC_DIR           = (_HERE / ".." / "static").resolve()

# Scan types that run the full tool suite (Anchore, Trivy, Checkov, etc.)
# Quick/stig/local_model scans do not produce complete vulnerability counts.
_COMPREHENSIVE_SCAN_TYPES = frozenset({"full", "nightly"})

# ── Audit logging ────────────────────────────────────────────
_AUDIT_LOG_FILE = (_HERE / ".." / "data" / "audit.log").resolve()


def _setup_audit_logger() -> logging.Logger:
    logger = logging.getLogger("epyon.audit")
    if not logger.handlers:
        _AUDIT_LOG_FILE.parent.mkdir(parents=True, exist_ok=True)
        handler = logging.FileHandler(str(_AUDIT_LOG_FILE), encoding="utf-8")
        handler.setFormatter(logging.Formatter(
            "%(asctime)s %(levelname)s %(message)s",
            datefmt="%Y-%m-%dT%H:%M:%SZ",
        ))
        logger.addHandler(handler)
        logger.setLevel(logging.INFO)
        logger.propagate = False
    return logger


_audit_logger = logging.getLogger("epyon.audit")


def _audit(request: Request, action: str, detail: str = "") -> None:
    """Write a single line to the audit log."""
    global _audit_logger
    if not _audit_logger.handlers:
        _audit_logger = _setup_audit_logger()
    client = getattr(request.client, "host", "unknown") if request.client else "unknown"
    _audit_logger.info("action=%s client=%s detail=%s", action, client, detail)

# ── Validation ────────────────────────────────────────────────
_SAFE_ID_RE      = re.compile(r"^[a-zA-Z0-9][a-zA-Z0-9_\-.]*$")
_JOB_ID_RE       = re.compile(r"^\d{14}$")
_APP_SCAN_RE     = re.compile(r"^(.+)_(\d{4}-\d{2}-\d{2})_(\d{2}-\d{2}-\d{2})$")
_VALID_SCAN_TYPES = {"quick", "full", "nightly", "baseline", "stig", "local_model"}
_TOKEN_RE        = re.compile(r"^(ghp_|github_pat_|ghs_|gho_)[a-zA-Z0-9_]+$")
_REPO_RE         = re.compile(r"^[a-zA-Z0-9_.\-]+/[a-zA-Z0-9_.\-]+$")

# ── Metrics cache ─────────────────────────────────────────────
_metrics_cache:    dict  = {}
_metrics_cache_ts: float = 0.0
_METRICS_TTL:      float = 300.0

# ── GitHub metrics cache ─────────────────────────────────────
_gh_metrics_cache:    dict  = {}
_gh_metrics_cache_ts: float = 0.0
_GH_METRICS_TTL:      float = 600.0  # 10 min — GH API is rate-limited
_GH_HISTORY_FILE = (_HERE / ".." / "data" / "github-signals-history.json").resolve()


def _load_gh_history() -> list[dict]:
    """Return the list of daily GitHub signals snapshots, or [] if none yet."""
    try:
        if _GH_HISTORY_FILE.exists():
            return json.loads(_GH_HISTORY_FILE.read_text(encoding="utf-8"))
    except Exception:
        pass
    return []


def _save_gh_snapshot(metrics: dict) -> None:
    """Append (or update) a daily aggregate snapshot of GitHub metrics to the history file."""
    try:
        dep   = metrics.get("dependabot", [])
        sec   = metrics.get("security_issues", [])
        runs  = metrics.get("workflow_runs", [])
        cov   = metrics.get("pr_scan_coverage", {})

        dep_open  = sum(r.get("open", 0)  for r in dep  if not r.get("error"))
        dep_fixed = sum(r.get("fixed", 0) for r in dep  if not r.get("error"))
        sec_open  = sum(r.get("open_security", 0) for r in sec if not r.get("error"))

        run_rates = [r["success_rate"] for r in runs if not r.get("error") and r.get("success_rate") is not None]
        workflow_success_rate = round(sum(run_rates) / len(run_rates)) if run_rates else None

        pr_scans  = sum(v.get("pr_scans", 0)     for v in cov.values())
        ci_scans  = sum(v.get("total_ci_scans", 0) for v in cov.values())
        pr_cov_pct = round(pr_scans / ci_scans * 100) if ci_scans > 0 else 0

        today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        snapshot = {
            "date":                   today,
            "ts":                     datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
            "dep_open":               dep_open,
            "dep_fixed":              dep_fixed,
            "sec_open":               sec_open,
            "workflow_success_rate":  workflow_success_rate,
            "pr_coverage_pct":        pr_cov_pct,
        }

        history = _load_gh_history()
        # Replace the entry for today if it exists; otherwise append
        history = [e for e in history if e.get("date") != today]
        history.append(snapshot)
        # Keep at most 365 daily entries
        history = sorted(history, key=lambda e: e["date"])[-365:]

        _GH_HISTORY_FILE.parent.mkdir(parents=True, exist_ok=True)
        _GH_HISTORY_FILE.write_text(json.dumps(history, indent=2), encoding="utf-8")
    except Exception:
        pass  # never let history writes break the main metrics response

# ── Scan data cache ───────────────────────────────────────────
# Short TTL caches for filesystem-heavy operations.
# find_scan_dirs() is called on every list endpoint; load_scan() reads 8+ files
# per scan directory. With 40+ scans these dominate response time.
_scan_cache:      dict[str, tuple[dict, float]] = {}  # scan_id → (data, monotonic_ts)
_dir_cache:       list | None = None
_dir_cache_ts:    float = 0.0
_SCAN_CACHE_TTL:  float = 30.0   # seconds — stale counts tolerable; scans take minutes
_DIR_CACHE_TTL:   float = 10.0   # seconds — new scans appear within 10 s


def _cached_find_scan_dirs() -> list:
    global _dir_cache, _dir_cache_ts
    now = time.monotonic()
    if _dir_cache is not None and (now - _dir_cache_ts) < _DIR_CACHE_TTL:
        return _dir_cache
    _dir_cache    = parsers.find_scan_dirs(EPYON_ROOT)
    _dir_cache_ts = now
    return _dir_cache


def _cached_load_scan(scan_dir) -> dict:
    scan_id = scan_dir.name
    now     = time.monotonic()
    cached  = _scan_cache.get(scan_id)
    if cached and (now - cached[1]) < _SCAN_CACHE_TTL:
        return cached[0]
    data = parsers.load_scan(scan_dir, EPYON_ROOT)
    _scan_cache[scan_id] = (data, now)
    return data


def _invalidate_scan_cache() -> None:
    """Call after a scan completes so the next request sees fresh data."""
    global _dir_cache, _dir_cache_ts
    _scan_cache.clear()
    _dir_cache    = None
    _dir_cache_ts = 0.0


async def _jira_post_scan(target_name: str) -> None:
    """Auto-reconcile Jira tickets after a scan finishes for target_name."""
    try:
        cfg = jira_client.read_config()
        if not cfg.get("auto_close") or not cfg.get("api_token"):
            return

        all_dirs = parsers.find_scan_dirs(EPYON_ROOT)
        target_dirs = sorted(
            [d for d in all_dirs if parsers.parse_dir_name(d.name)["target"] == target_name],
            key=lambda d: d.name,
            reverse=True,
        )
        if len(target_dirs) < 2:
            return  # need at least two scans to compare

        current_raw  = (parsers.load_enriched_findings(target_dirs[0])
                        or parsers.parse_scan_findings(target_dirs[0]))
        previous_raw = (parsers.load_enriched_findings(target_dirs[1])
                        or parsers.parse_scan_findings(target_dirs[1]))

        await jira_client.reconcile_and_save(
            target_name,
            jira_client.flatten_findings(current_raw),
            jira_client.flatten_findings(previous_raw),
            cfg,
        )
    except Exception:
        pass  # never let Jira errors break the scan pipeline


def _on_scan_complete(target_name: str = "", scan_name: str = "") -> None:
    """Callback invoked by jobs.py when a scan finishes."""
    _invalidate_scan_cache()
    if target_name:
        asyncio.create_task(_jira_post_scan(target_name))


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _read_github_config() -> dict:
    try:
        return json.loads(GITHUB_CONFIG_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _write_github_config(cfg: dict) -> None:
    GITHUB_CONFIG_FILE.write_text(json.dumps(cfg, indent=2), encoding="utf-8")


# ── Lifespan ─────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(_: FastAPI):
    async def _auto_sync_loop() -> None:
        await asyncio.sleep(10)
        while True:
            cfg = _read_github_config()
            if cfg.get("token") and cfg.get("repos"):
                await github_sync.trigger_sync(
                    GITHUB_CONFIG_FILE, EPYON_ROOT, parsers.find_scan_dirs,
                    on_complete=_invalidate_scan_cache,
                )
            await asyncio.sleep(15 * 60)  # re-check every 15 minutes

    asyncio.create_task(_auto_sync_loop())
    yield


# ── Security header constants ─────────────────────────────────
# CSP allows 'unsafe-inline' because app.js builds HTML via template literals
# with inline onclick handlers.  Removing it requires a full JS refactor and
# is tracked as a future hardening item (see APSC-DV-001600 evidence).
_CSP = (
    "default-src 'self'; "
    "script-src 'self' 'unsafe-inline'; "
    "style-src 'self' 'unsafe-inline'; "
    "img-src 'self' data: blob:; "
    "connect-src 'self'; "
    "font-src 'self'; "
    "object-src 'none'; "
    "base-uri 'self'; "
    "form-action 'self';"
)
_PERMISSIONS_POLICY = "geolocation=(), microphone=(), camera=(), payment=(), usb=()"


# ── App ───────────────────────────────────────────────────────

app = FastAPI(title="Epyon Web", lifespan=lifespan)

# Restrict CORS to same-origin by default; override via EPYON_ALLOWED_ORIGINS
# (comma-separated list).  Never use wildcard in production deployments.
_ALLOWED_ORIGINS = [
    o.strip()
    for o in os.getenv(
        "EPYON_ALLOWED_ORIGINS",
        "http://localhost:8000,http://127.0.0.1:8000",
    ).split(",")
    if o.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization", "X-Requested-With"],
)


class _SecurityHeadersMiddleware(BaseHTTPMiddleware):
    """Apply security headers to every response (APSC-DV-001600 et al.)."""

    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"]  = "nosniff"
        response.headers["X-Frame-Options"]         = "SAMEORIGIN"
        response.headers["X-XSS-Protection"]        = "0"
        response.headers["Referrer-Policy"]         = "strict-origin-when-cross-origin"
        response.headers["Permissions-Policy"]      = _PERMISSIONS_POLICY
        response.headers["Content-Security-Policy"] = _CSP
        if request.url.path.startswith("/api/"):
            response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, private"
            response.headers["Pragma"]        = "no-cache"
        return response


app.add_middleware(_SecurityHeadersMiddleware)


@app.exception_handler(Exception)
async def _unhandled_exception_handler(request: Request, exc: Exception) -> JSONResponse:
    """Return a generic 500 without leaking internal details (APSC-DV-002390)."""
    _audit(request, "server_error", type(exc).__name__)
    return JSONResponse(status_code=500, content={"detail": "An internal error occurred."})


def _sec_headers(response: Response) -> None:
    response.headers["X-Content-Type-Options"]  = "nosniff"
    response.headers["X-Frame-Options"]         = "SAMEORIGIN"
    response.headers["X-XSS-Protection"]        = "0"
    response.headers["Referrer-Policy"]         = "strict-origin-when-cross-origin"
    response.headers["Permissions-Policy"]      = _PERMISSIONS_POLICY
    response.headers["Content-Security-Policy"] = _CSP
    response.headers["Cache-Control"]           = "no-store, no-cache, must-revalidate, private"
    response.headers["Pragma"]                  = "no-cache"


# ── Health ────────────────────────────────────────────────────

@app.get("/api/health")
def health(response: Response):
    _sec_headers(response)
    return {"status": "ok", "epyon_root": str(EPYON_ROOT), "version": EPYON_VERSION}


# ── Stats ─────────────────────────────────────────────────────

@app.get("/api/stats")
def stats(response: Response):
    _sec_headers(response)
    hidden = _load_hidden_apps()
    scans = [_cached_load_scan(d) for d in _cached_find_scan_dirs()]
    by_target: dict[str, dict] = {}
    for s in scans:
        t = s["target"]
        if t in hidden:
            continue
        if t not in by_target or s.get("timestamp", "") > by_target[t].get("timestamp", ""):
            by_target[t] = s
    latest = list(by_target.values())
    return {
        "total_applications": len(by_target),
        "total_scans":        len(scans),
        "critical": sum(s["critical"] for s in latest),
        "high":     sum(s["high"]     for s in latest),
        "medium":   sum(s["medium"]   for s in latest),
        "low":      sum(s["low"]      for s in latest),
    }


# ── Scan history ──────────────────────────────────────────────

@app.get("/api/scan-history")
def scan_history(response: Response):
    _sec_headers(response)
    hidden  = _load_hidden_apps()
    scans   = [_cached_load_scan(d) for d in _cached_find_scan_dirs()
               if parsers.parse_dir_name(d.name)["target"] not in hidden]
    targets = sorted({s["target"] for s in scans})
    users   = sorted({s["user"] for s in scans if s.get("user")})
    return {
        "generated_at": _now(),
        "total_scans":  len(scans),
        "targets":      targets,
        "users":        users,
    }


# ── Applications ──────────────────────────────────────────────

def _load_hidden_apps() -> set[str]:
    try:
        if HIDDEN_APPS_FILE.exists():
            return set(json.loads(HIDDEN_APPS_FILE.read_text()))
    except Exception:
        pass
    return set()


def _save_hidden_apps(hidden: set[str]) -> None:
    HIDDEN_APPS_FILE.parent.mkdir(parents=True, exist_ok=True)
    HIDDEN_APPS_FILE.write_text(json.dumps(sorted(hidden), indent=2))


def _load_monitored_apps() -> set[str]:
    try:
        if MONITORED_APPS_FILE.exists():
            return set(json.loads(MONITORED_APPS_FILE.read_text()))
    except Exception:
        pass
    return set()


def _save_monitored_apps(monitored: set[str]) -> None:
    MONITORED_APPS_FILE.parent.mkdir(parents=True, exist_ok=True)
    MONITORED_APPS_FILE.write_text(json.dumps(sorted(monitored), indent=2))


def _load_registered_apps() -> list[dict]:
    try:
        if REGISTERED_APPS_FILE.exists():
            return json.loads(REGISTERED_APPS_FILE.read_text(encoding="utf-8"))
    except Exception:
        pass
    return []


def _save_registered_apps(apps: list[dict]) -> None:
    REGISTERED_APPS_FILE.parent.mkdir(parents=True, exist_ok=True)
    REGISTERED_APPS_FILE.write_text(json.dumps(apps, indent=2), encoding="utf-8")


@app.post("/api/applications", status_code=201)
async def register_application(request: Request, response: Response):
    """Register a new application by name and URL (no scan triggered)."""
    _sec_headers(response)
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "Invalid JSON")
    name = (body.get("name") or "").strip()
    url  = (body.get("url")  or "").strip()
    if not name:
        raise HTTPException(400, "name is required")
    if not _SAFE_ID_RE.match(name):
        raise HTTPException(400, "name must be alphanumeric with hyphens/underscores/dots only")
    valid_prefixes = ["https://", "http://", "git@"]
    if url and not any(url.startswith(p) for p in valid_prefixes):
        raise HTTPException(400, "url must be an HTTPS or SSH Git URL")
    if url and re.search(r"[;&|`$\(\)\n\r<>'\"]", url):
        raise HTTPException(400, "url contains invalid characters")
    apps = _load_registered_apps()
    apps = [a for a in apps if a["name"] != name]
    apps.append({"name": name, "url": url, "added_at": _now()})
    _save_registered_apps(apps)
    return {"registered": name, "url": url}


@app.get("/api/applications")
def applications(response: Response):
    _sec_headers(response)
    hidden   = _load_hidden_apps()
    monitored = _load_monitored_apps()
    scans = [_cached_load_scan(d) for d in _cached_find_scan_dirs()]
    by_target: dict[str, list] = {}
    for s in scans:
        if s["target"] not in hidden:
            by_target.setdefault(s["target"], []).append(s)

    result = []
    for name, tscans in by_target.items():
        tscans.sort(key=lambda x: x.get("timestamp", ""), reverse=True)
        latest = tscans[0] if tscans else {}

        # For vulnerability counts, prefer the latest comprehensive scan (full/nightly)
        # so that quick/stig/local_model scans don't zero out the displayed counts.
        latest_comprehensive = next(
            (s for s in tscans if s.get("scan_type") in _COMPREHENSIVE_SCAN_TYPES),
            latest,
        )

        # Find the latest scan that actually produced STIG output
        latest_stig = next((s for s in tscans if s.get("stig_total", 0) > 0), None)

        result.append({
            "name":                      name,
            "scan_count":                len(tscans),
            "last_scanned":              latest.get("timestamp", ""),
            "scan_type":                 latest.get("scan_type", ""),
            "critical":                  latest_comprehensive.get("critical", 0),
            "high":                      latest_comprehensive.get("high", 0),
            "medium":                    latest_comprehensive.get("medium", 0),
            "low":                       latest_comprehensive.get("low", 0),
            "status":                    parsers.get_status(latest_comprehensive),
            "latest_scan_id":            latest.get("scan_id", ""),
            "comprehensive_scan_id":     latest_comprehensive.get("scan_id", ""),
            "comprehensive_timestamp":   latest_comprehensive.get("timestamp", ""),
            "comprehensive_scan_type":   latest_comprehensive.get("scan_type", ""),
            "stig_total":                latest_stig.get("stig_total", 0)          if latest_stig else 0,
            "stig_open":           latest_stig.get("stig_open", 0)           if latest_stig else 0,
            "stig_pass":           latest_stig.get("stig_pass", 0)           if latest_stig else 0,
            "stig_na":             latest_stig.get("stig_na", 0)             if latest_stig else 0,
            "latest_stig_scan_id": latest_stig.get("scan_id", "")            if latest_stig else "",
            "has_stig_report":     latest_stig.get("has_stig_report", False) if latest_stig else False,
            "has_stig_cklb":       latest_stig.get("has_stig_cklb", False)   if latest_stig else False,
            "url":                 "",
            "monitored":           name in monitored,
        })

    # Merge in registered-but-unscanned apps
    scanned_names = {r["name"] for r in result}
    registered = _load_registered_apps()
    for reg in registered:
        rname = reg.get("name")
        if not rname:
            continue
        if rname in hidden or rname in scanned_names:
            # If it has been scanned, backfill the URL onto the existing entry
            for r in result:
                if r["name"] == rname and not r["url"]:
                    r["url"] = reg.get("url", "")
            continue
        result.append({
            "name":                rname,
            "scan_count":          0,
            "last_scanned":        "",
            "scan_type":           "",
            "critical":            0,
            "high":                0,
            "medium":              0,
            "low":                 0,
            "status":              "unknown",
            "latest_scan_id":      "",
            "stig_total":          0,
            "stig_open":           0,
            "stig_pass":           0,
            "stig_na":             0,
            "latest_stig_scan_id": "",
            "has_stig_report":     False,
            "has_stig_cklb":       False,
            "url":                 reg.get("url", ""),
            "added_at":            reg.get("added_at", ""),
            "monitored":           rname in monitored,
        })

    result.sort(key=lambda x: x.get("last_scanned", "") or x.get("added_at", ""), reverse=True)
    return result


# ── Application action endpoints ─────────────────────────────────────────────
# Names are passed in the request body (POST) or as a ?name= query param
# (DELETE) to avoid Starlette normalising %2F → / in URL path parameters,
# which breaks routing when the application name is a full Git URL.

def _require_app_name(name: str | None) -> str:
    name = (name or "").strip()
    if not name:
        raise HTTPException(400, "name is required")
    return name


@app.post("/api/applications/hide")
async def hide_application(request: Request, response: Response):
    _sec_headers(response)
    body = await request.json()
    name = _require_app_name(body.get("name"))
    hidden = _load_hidden_apps()
    hidden.add(name)
    _save_hidden_apps(hidden)
    return {"hidden": name}


@app.post("/api/applications/restore")
async def restore_application(request: Request, response: Response):
    _sec_headers(response)
    body = await request.json()
    name = _require_app_name(body.get("name"))
    hidden = _load_hidden_apps()
    hidden.discard(name)
    _save_hidden_apps(hidden)
    return {"restored": name}


@app.post("/api/applications/monitored")
async def set_monitored(request: Request, response: Response):
    _sec_headers(response)
    body = await request.json()
    name = _require_app_name(body.get("name"))
    monitored = _load_monitored_apps()
    monitored.add(name)
    _save_monitored_apps(monitored)
    return {"monitored": name}


@app.delete("/api/applications/monitored")
def unset_monitored(name: str, response: Response):
    _sec_headers(response)
    name = _require_app_name(name)
    monitored = _load_monitored_apps()
    monitored.discard(name)
    _save_monitored_apps(monitored)
    return {"unmonitored": name}


@app.delete("/api/applications/data")
def delete_application(name: str, request: Request, response: Response):
    """Permanently delete all scan directories for an application."""
    _sec_headers(response)
    name = _require_app_name(name)
    _audit(request, "delete_application", f"app={name}")
    scan_dirs = [
        d for d in parsers.find_scan_dirs(EPYON_ROOT)
        if parsers.parse_dir_name(d.name)["target"] == name
    ]
    deleted = []
    for d in scan_dirs:
        try:
            d.resolve().relative_to(EPYON_ROOT.resolve())
        except ValueError:
            raise HTTPException(403, "Access denied")
        shutil.rmtree(d)
        deleted.append(d.name)
    hidden = _load_hidden_apps()
    if name in hidden:
        hidden.discard(name)
        _save_hidden_apps(hidden)
    return {"deleted": deleted}


@app.delete("/api/scans/{scan_id}")
def delete_scan(scan_id: str, request: Request, response: Response):
    """Permanently delete a single scan directory."""
    _sec_headers(response)
    _audit(request, "delete_scan", f"scan_id={scan_id}")
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")
    try:
        matched.resolve().relative_to(EPYON_ROOT.resolve())
    except ValueError:
        raise HTTPException(403, "Access denied")
    shutil.rmtree(matched)
    return {"deleted": scan_id}


@app.get("/api/applications-hidden")
def hidden_applications(response: Response):
    _sec_headers(response)
    return sorted(_load_hidden_apps())


@app.get("/api/applications/{name}/scans")
def app_scans(name: str, response: Response):
    _sec_headers(response)
    if not _SAFE_ID_RE.match(name):
        raise HTTPException(400, "Invalid application name")
    scans = [
        parsers.load_scan(d, EPYON_ROOT)
        for d in parsers.find_scan_dirs(EPYON_ROOT)
        if parsers.parse_dir_name(d.name)["target"] == name
    ]
    scans.sort(key=lambda x: x.get("timestamp", ""), reverse=True)
    return scans


@app.post("/api/applications/{name}/isso-summary")
async def app_isso_summary(name: str, response: Response):
    """Per-application ISSO compliance brief: NIST control mapping, STIG, POA&M, accepted risk."""
    _sec_headers(response)
    if not _SAFE_ID_RE.match(name):
        raise HTTPException(400, "Invalid application name")

    all_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    target_entries: list[tuple] = []
    for d in all_dirs:
        meta = parsers.parse_dir_name(d.name)
        if meta["target"] == name:
            target_entries.append((d, meta["timestamp"]))

    if not target_entries:
        raise HTTPException(404, f"No scans found for application '{name}'")

    # Sort newest first
    target_entries.sort(key=lambda x: x[1], reverse=True)

    # Best comprehensive scan dir for vuln findings
    def _best_scan_dir(entries: list[tuple]) -> object:
        for d, _ in entries:
            st = (parsers._read_json(d / "scan-metadata.json") or {}).get("scan_type", "")
            if st in _COMPREHENSIVE_SCAN_TYPES:
                return d
        return entries[0][0]

    scan_dir  = _best_scan_dir(target_entries)
    scan_meta = parsers.load_scan(scan_dir, EPYON_ROOT)
    # Prefer the pre-built deduplicated summary; fall back to raw parse only if absent
    findings  = parsers.load_enriched_findings(scan_dir) or parsers.parse_scan_findings(scan_dir)

    # STIG: walk newest-first until we find a scan with stig-results-*.json files;
    # merge all controls from that scan into a single flat list.
    stig_detail: list[dict] | None = None
    for candidate, _ in target_entries:
        stig_files = sorted(candidate.glob("stig-results-*.json"))
        if not stig_files:
            continue
        merged: list[dict] = []
        for sf in stig_files:
            try:
                raw = json.loads(sf.read_text(encoding="utf-8"))
                results = raw.get("assessments", raw) if "assessments" in raw else raw
                # Load matching controls file
                slug = sf.stem[len("stig-results-"):]
                cf = candidate / f"stig-controls-{slug}.json"
                controls_data: dict = {}
                try:
                    controls_data = json.loads(cf.read_text(encoding="utf-8"))
                except Exception:
                    pass
                for c in controls_data.get("controls", []):
                    vid = c.get("vuln_id", "")
                    assessed = results.get(vid, {})
                    merged.append({
                        "vuln_id":  vid,
                        "title":    c.get("title", ""),
                        "severity": c.get("severity", ""),
                        "status":   assessed.get("status", "Not Reviewed"),
                    })
            except Exception:
                pass
        if merged:
            stig_detail = merged
            break

    # Suppressions: walk newest-first until non-empty
    suppressions: list[dict] = []
    for candidate, _ in target_entries:
        batch = parsers.parse_suppressed_findings(candidate)
        if batch:
            suppressions = batch
            break

    # Scan history: last 8 scans, summary counts only
    scan_history = [
        {
            "timestamp": m.get("timestamp", ""),
            "scan_type": m.get("scan_type", ""),
            "critical":  m.get("critical", 0),
            "high":      m.get("high", 0),
            "medium":    m.get("medium", 0),
            "low":       m.get("low", 0),
        }
        for m in (
            parsers.load_scan(d, EPYON_ROOT)
            for d, _ in target_entries[:8]
        )
    ]

    _metrics = _build_summary_metrics()

    try:
        summary = await openai_summary.generate_app_isso_summary(
            app_name=name,
            scan_meta=scan_meta,
            findings=findings,
            stig_detail=stig_detail,
            suppressions=suppressions,
            scan_history=scan_history,
            metrics=_metrics,
        )
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))
    except Exception as exc:
        raise HTTPException(502, f"OpenAI request failed: {exc}")

    return {"app": name, "summary": summary}


# ── Scan detail ───────────────────────────────────────────────

@app.get("/api/scans/{scan_id}/dashboard")
def scan_dashboard(scan_id: str, response: Response):
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")
    dashboard = matched / "consolidated-reports" / "dashboards" / "security-dashboard.html"
    try:
        dashboard.resolve().relative_to(EPYON_ROOT.resolve())
    except ValueError:
        raise HTTPException(403, "Access denied")
    if not dashboard.exists():
        raise HTTPException(404, "Dashboard not generated for this scan")
    return FileResponse(str(dashboard), media_type="text/html")


@app.get("/api/scans/{scan_id}/stig-findings-md")
def stig_findings_md(scan_id: str, response: Response):
    """Serve the primary STIG findings markdown for a given scan."""
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")
    app_slug = re.sub(r"[^a-z0-9]+", "-", re.sub(r"_\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2}$", "", scan_id).lower()).strip("-")
    report = matched / f"findings-{app_slug}.md"
    try:
        report.resolve().relative_to(EPYON_ROOT.resolve())
    except ValueError:
        raise HTTPException(403, "Access denied")
    if not report.exists():
        raise HTTPException(404, "STIG findings not generated for this scan")
    return FileResponse(str(report), media_type="text/markdown",
                        headers={"Content-Disposition": f'attachment; filename="findings-{scan_id}.md"'})


@app.get("/api/scans/{scan_id}/stig-findings-cklb")
def stig_findings_cklb(scan_id: str, response: Response):
    """Serve the primary STIG findings CKLB for a given scan."""
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")
    app_slug = re.sub(r"[^a-z0-9]+", "-", re.sub(r"_\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2}$", "", scan_id).lower()).strip("-")
    report = matched / f"findings-{app_slug}.cklb"
    try:
        report.resolve().relative_to(EPYON_ROOT.resolve())
    except ValueError:
        raise HTTPException(403, "Access denied")
    if not report.exists():
        raise HTTPException(404, "STIG findings CKLB not generated for this scan")
    return FileResponse(str(report), media_type="application/json",
                        headers={"Content-Disposition": f'attachment; filename="findings-{scan_id}.cklb"'})


@app.get("/api/scans/{scan_id}/stig-data")
def stig_data(scan_id: str, response: Response):
    """Return merged controls + assessment results for all STIGs in a scan as JSON."""
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")

    stigs: list[dict] = []
    for results_file in sorted(matched.glob("stig-results-*.json")):
        slug = results_file.stem[len("stig-results-"):]
        controls_file = matched / f"stig-controls-{slug}.json"

        results: dict = {}
        try:
            raw = json.loads(results_file.read_text(encoding="utf-8"))
            # Support new wrapped format {assessments: {...}, token_usage: {...}}
            # and old flat format {vuln_id: {status, evidence}, ...}
            results = raw.get("assessments", raw) if "assessments" in raw else raw
        except Exception:
            pass

        controls_data: dict = {}
        try:
            controls_data = json.loads(controls_file.read_text(encoding="utf-8"))
        except Exception:
            pass

        controls_list = controls_data.get("controls", [])
        merged: list[dict] = []
        for c in controls_list:
            vid = c.get("vuln_id", "")
            assessed = results.get(vid, {})
            merged.append({
                "vuln_id":        vid,
                "group_id":       c.get("group_id", ""),
                "rule_id":        c.get("rule_id", ""),
                "number":         c.get("number"),
                "severity":       c.get("severity", ""),
                "title":          c.get("title", ""),
                "check_content":  c.get("check_content", ""),
                "fix_text":       c.get("fix_text", ""),
                "discussion":     c.get("discussion", ""),
                "status":         assessed.get("status",          "Not Reviewed"),
                "evidence":       assessed.get("evidence",         ""),
                "confidence":     assessed.get("confidence",       0),
                "locked_by_human": bool(assessed.get("locked_by_human", False)),
            })

        app_slug  = re.sub(r"[^a-z0-9]+", "-", re.sub(r"_\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2}$", "", scan_id).lower()).strip("-")
        md_file   = matched / f"findings-{app_slug}-{slug}.md"
        cklb_file = matched / f"findings-{app_slug}-{slug}.cklb"
        stigs.append({
            "slug":        slug,
            "stig_name":   controls_data.get("stig_name", slug),
            "release_info": controls_data.get("release_info", ""),
            "total":       len(merged),
            "controls":    merged,
            "md_url":      f"/api/scans/{scan_id}/stig-findings/{app_slug}-{slug}.md"   if md_file.exists()   else None,
            "cklb_url":    f"/api/scans/{scan_id}/stig-findings/{app_slug}-{slug}.cklb" if cklb_file.exists() else None,
        })

    return {"scan_id": scan_id, "stigs": stigs}


_SAFE_SLUG_RE = re.compile(r"^[a-zA-Z0-9][a-zA-Z0-9_\-]*\.(md|cklb)$")
_VALID_STIG_STATUSES = {"Not a Finding", "Not Applicable", "Open", "Not Reviewed"}


class StigOverride(BaseModel):
    slug: str
    vuln_id: str
    status: str
    evidence: str = ""
    justification: str = ""  # human rationale stored alongside evidence
    clear_lock: bool = False  # when True, removes the human lock and resets to Open


@app.patch("/api/scans/{scan_id}/stig-override")
def stig_override(scan_id: str, body: StigOverride, response: Response):
    """Apply a human-reviewed override to a single STIG control.

    Sets status, updates evidence, marks locked_by_human=True and
    confidence=95 so the freeze logic carries this forward on subsequent scans.
    The justification is prepended to the evidence text so it is visible in
    the findings report.
    """
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    # Validate slug to prevent path traversal
    if not re.match(r'^[a-z0-9][a-z0-9\-]*$', body.slug):
        raise HTTPException(400, "Invalid slug")
    if body.status not in _VALID_STIG_STATUSES:
        raise HTTPException(400, f"Invalid status. Must be one of: {sorted(_VALID_STIG_STATUSES)}")
    # Validate vuln_id format (e.g. APSC-DV-000070, TCAT-AS-000010)
    if not re.match(r'^[A-Z0-9]{2,}-[A-Z0-9]{2,}-[0-9]{4,}$', body.vuln_id):
        raise HTTPException(400, "Invalid vuln_id format")

    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")

    results_file = matched / f"stig-results-{body.slug}.json"
    if not results_file.exists():
        raise HTTPException(404, f"STIG results file not found for slug: {body.slug}")

    # Resolve and verify the path is inside the scan directory (no traversal)
    try:
        results_file.resolve().relative_to(EPYON_ROOT.resolve())
    except ValueError:
        raise HTTPException(403, "Access denied")

    try:
        raw = json.loads(results_file.read_text(encoding="utf-8"))
    except Exception as exc:
        raise HTTPException(500, f"Could not read results file: {exc}") from exc

    assessments = raw.get("assessments", raw) if "assessments" in raw else raw

    # Build the updated evidence text — prepend justification if provided
    new_evidence = body.evidence
    if body.justification:
        prefix = f"[Human override] {body.justification}\n"
        if not new_evidence.startswith(prefix):
            new_evidence = prefix + new_evidence

    if body.clear_lock:
        # Remove the human lock; reset to Open with confidence 0 so AI re-assesses next scan
        assessments[body.vuln_id] = {
            **assessments.get(body.vuln_id, {}),
            "status":          "Open",
            "evidence":        "",
            "confidence":      0,
            "locked_by_human": False,
        }
    else:
        assessments[body.vuln_id] = {
            **assessments.get(body.vuln_id, {}),
            "status":          body.status,
            "evidence":        new_evidence,
            "confidence":      95,
            "locked_by_human": True,
        }

    # Write back in the wrapped format
    if "assessments" in raw:
        raw["assessments"] = assessments
    else:
        raw = {"assessments": assessments}

    results_file.write_text(json.dumps(raw, indent=2, ensure_ascii=False), encoding="utf-8")
    results_file.chmod(0o600)

    return {
        "scan_id":  scan_id,
        "slug":     body.slug,
        "vuln_id":  body.vuln_id,
        "status":   "Open" if body.clear_lock else body.status,
        "locked":   not body.clear_lock,
    }


@app.get("/api/scans/{scan_id}/stig-findings/{filename}")
def stig_findings_named(scan_id: str, filename: str, response: Response):
    """Serve a per-STIG named findings file (findings-{slug}.md or .cklb)."""
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    if not _SAFE_SLUG_RE.match(filename):
        raise HTTPException(400, "Invalid filename")
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")
    report = matched / f"findings-{filename}"
    try:
        report.resolve().relative_to(EPYON_ROOT.resolve())
    except ValueError:
        raise HTTPException(403, "Access denied")
    if not report.exists():
        raise HTTPException(404, f"File not found: findings-{filename}")
    ext = filename.rsplit(".", 1)[-1]
    media = "text/markdown" if ext == "md" else "application/json"
    return FileResponse(str(report), media_type=media,
                        headers={"Content-Disposition": f'attachment; filename="findings-{filename}"'})


@app.get("/api/scans/{scan_id}")
def scan_detail(scan_id: str, response: Response):
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")
    data = parsers.load_scan(matched, EPYON_ROOT)
    data["findings"] = parsers.load_enriched_findings(matched) or parsers.parse_scan_findings(matched)
    data["sbom"] = parsers.load_sbom_packages(matched)
    data["api_discovery"] = parsers.load_api_discovery(matched)
    return data


@app.get("/api/scans/{scan_id}/sbom")
def scan_sbom(scan_id: str, response: Response):
    """Return SBOM package list for a scan (reads syft-json directly)."""
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")
    return parsers.load_sbom_packages(matched)


@app.get("/api/scans/{scan_id}/sbom/cyclonedx")
def scan_sbom_cyclonedx(scan_id: str, response: Response):
    """Download the CycloneDX SBOM JSON for a scan."""
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")
    sbom_dir = matched / "sbom"
    # Prefer the dedicated cyclonedx file, fall back to any *.cyclonedx.json
    candidates = [
        sbom_dir / "filesystem.cyclonedx.json",
        *sorted(sbom_dir.glob("*.cyclonedx.json")),
    ]
    cdx_file = next((f for f in candidates if f.exists()), None)
    if cdx_file is None:
        raise HTTPException(404, "CycloneDX SBOM not found for this scan")
    filename = f"sbom-{scan_id}.cyclonedx.json"
    return FileResponse(
        str(cdx_file),
        media_type="application/vnd.cyclonedx+json",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/scans/{scan_id}/download")
def scan_download_zip(scan_id: str, response: Response):
    """Stream the entire scan directory as a ZIP archive for ATO/IATT submissions."""
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")

    epyon_root_resolved = EPYON_ROOT.resolve()

    def _generate_zip():
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
            for file_path in sorted(matched.rglob("*")):
                if not file_path.is_file():
                    continue
                # Skip large/noisy files that add no ATO value
                rel = file_path.relative_to(matched)
                rel_str = str(rel)
                if any(part.startswith(".") for part in rel.parts):
                    continue  # skip hidden dirs (.scannerwork etc)
                if file_path.suffix in (".tar", ".gz", ".db", ".zip"):
                    continue
                # Path traversal guard
                try:
                    file_path.resolve().relative_to(epyon_root_resolved)
                except ValueError:
                    continue
                zf.write(file_path, arcname=str(Path(scan_id) / rel))
        buf.seek(0)
        yield from buf

    filename = f"epyon-scan-{scan_id}.zip"
    return StreamingResponse(
        _generate_zip(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Trigger scan ──────────────────────────────────────────────

@app.post("/api/scans", status_code=202)
async def trigger_scan(request: Request, response: Response):
    _sec_headers(response)
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "Invalid JSON")

    target    = (body.get("target") or "").strip()
    scan_type = body.get("scan_type", "full")
    run_garak = bool(body.get("run_garak", False))
    # run_stig is implicit when scan_type == "stig"; can also be set explicitly
    # for full/nightly scans that should include the STIG layer.
    run_stig  = bool(body.get("run_stig", False)) or scan_type == "stig"

    if not target:
        raise HTTPException(400, "target is required")
    valid_prefixes = ["/", "./", "../", "https://", "http://", "git@"]
    if not any(target.startswith(p) for p in valid_prefixes):
        raise HTTPException(400, "target must be an absolute path, relative path, or Git URL")
    if re.search(r"[;&|`$\(\)\n\r<>]", target):
        raise HTTPException(400, "target contains invalid characters")
    if scan_type not in _VALID_SCAN_TYPES:
        raise HTTPException(400, f"scan_type must be one of: {sorted(_VALID_SCAN_TYPES)}")

    script_path = SCRIPTS_DIR / "run-epyon-scan-ci.sh"
    if not script_path.exists():
        raise HTTPException(500, "Scan script not found")

    job_id = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    _audit(request, "scan_triggered", f"target={target} scan_type={scan_type}")
    job = job_store.create_job(job_id, target, scan_type)
    job_store._on_scan_complete_cb = _on_scan_complete
    asyncio.create_task(
        job_store.run_scan_job(job_id, target, scan_type, script_path, EPYON_ROOT,
                               run_garak=run_garak, run_stig=run_stig)
    )
    return {"job_id": job_id, "status": "queued"}


# ── Jobs ──────────────────────────────────────────────────────

@app.get("/api/jobs")
def list_jobs(response: Response):
    _sec_headers(response)
    return sorted(job_store.jobs.values(), key=lambda j: j.get("started_at", ""), reverse=True)


@app.get("/api/jobs/{job_id}")
def get_job(job_id: str, response: Response):
    _sec_headers(response)
    if not _JOB_ID_RE.match(job_id):
        raise HTTPException(400, "Invalid job_id")
    job = job_store.jobs.get(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    return job


@app.post("/api/jobs/{job_id}/cancel")
def cancel_job(job_id: str, response: Response):
    _sec_headers(response)
    if not _JOB_ID_RE.match(job_id):
        raise HTTPException(400, "Invalid job_id")
    if job_id not in job_store.jobs:
        raise HTTPException(404, "Job not found")
    if job_store.jobs[job_id]["status"] not in ("queued", "running"):
        raise HTTPException(409, "Job is not running")
    job_store.cancel_job(job_id)
    return {"job_id": job_id, "status": "cancelled"}


# ── Metrics ───────────────────────────────────────────────────

@app.get("/api/metrics")
def get_metrics(response: Response):
    _sec_headers(response)
    global _metrics_cache, _metrics_cache_ts
    now = time.monotonic()
    if _metrics_cache and (now - _metrics_cache_ts) < _METRICS_TTL:
        return _metrics_cache

    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    all_scans = [parsers.load_scan(d, EPYON_ROOT) for d in scan_dirs]

    trend = sorted(
        [
            {
                "scan_id":   s["scan_id"],
                "target":    s["target"],
                "timestamp": s["timestamp"],
                "critical":  s["critical"],
                "high":      s["high"],
                "medium":    s["medium"],
                "low":       s["low"],
            }
            for s in all_scans if s.get("timestamp")
        ],
        key=lambda x: x["timestamp"],
    )[-60:]

    by_target: dict[str, list] = {}
    for s, d in zip(all_scans, scan_dirs):
        by_target.setdefault(s["target"], []).append((s, d))

    # ── Monitored-app filter ──────────────────────────────────
    monitored = _load_monitored_apps()
    if monitored:
        by_target = {k: v for k, v in by_target.items() if k in monitored}
        trend = [t for t in trend if t["target"] in monitored]
    metrics_filtered = bool(monitored)
    monitored_count  = len(monitored)

    tool_counts:      dict[str, dict] = {}
    tool_app_counts:  dict[str, dict[str, int]] = {}
    total_with_fix    = 0
    total_without_fix = 0
    total_suppressed_instances = 0
    cve_counts:       dict[str, dict] = {}

    for _target, scan_list in by_target.items():
        scan_list.sort(key=lambda x: x[0].get("timestamp", ""), reverse=True)
        _, latest_dir = scan_list[0]
        total_suppressed_instances += parsers.count_suppressed_instances(latest_dir)
        findings = parsers.parse_scan_findings(latest_dir)
        if not findings:
            continue
        for sev in ("critical", "high", "medium", "low"):
            for f in findings.get(f"{sev}_findings", []):
                tool = f.get("tool") or "Unknown"
                if tool not in tool_counts:
                    tool_counts[tool] = {"critical": 0, "high": 0, "medium": 0, "low": 0, "total": 0}
                tool_counts[tool][sev]    += 1
                tool_counts[tool]["total"] += 1
                tool_app_counts.setdefault(tool, {})
                tool_app_counts[tool][_target] = tool_app_counts[tool].get(_target, 0) + 1
                if f.get("fixed_version"):
                    total_with_fix += 1
                else:
                    total_without_fix += 1
                cve_id = f.get("id") or ""
                if cve_id.startswith("CVE-"):
                    ftype = (f.get("type") or "").lower()
                    ftool = (f.get("tool") or "").lower()
                    if ftype == "container_vulnerability" or "anchore" in ftool:
                        cve_source = "container"
                    else:
                        cve_source = "code"
                    if cve_id not in cve_counts:
                        cve_counts[cve_id] = {
                            "count":    0,
                            "severity": sev,
                            "title":    (f.get("title") or "")[:120],
                            "apps":     set(),
                            "sources":  set(),
                        }
                    cve_counts[cve_id]["count"] += 1
                    cve_counts[cve_id]["apps"].add(_target)
                    cve_counts[cve_id]["sources"].add(cve_source)

    top_cves = sorted(cve_counts.items(), key=lambda x: x[1]["count"], reverse=True)[:20]

    scan_frequency = {
        target: {
            "total": len(scan_list),
            "dates": sorted(
                s.get("timestamp", "")[:10]
                for s, _ in scan_list if s.get("timestamp")
            ),
        }
        for target, scan_list in by_target.items()
    }

    # ── Merges to main ─────────────────────────────────────────
    merge_dates_by_target: dict[str, list[str]] = {}
    for _target, scan_list in by_target.items():
        for s, _ in scan_list:
            ci = s.get("ci_source", {})
            if ci.get("event") in ("push", "workflow_dispatch") and ci.get("branch") in ("main", "master", "dev"):
                date = (s.get("timestamp") or "")[:10]
                if date:
                    merge_dates_by_target.setdefault(_target, []).append(date)

    merges_to_main = {
        target: {
            "total": len(dates),
            "dates": sorted(dates),
        }
        for target, dates in merge_dates_by_target.items()
    }
    total_merges_to_main = sum(len(d) for d in merge_dates_by_target.values())

    # ── SLA Compliance Rate ────────────────────────────────────
    # SLA thresholds (days): critical=7, high=30, medium=90, low=180
    SLA_DAYS = {"critical": 7, "high": 30, "medium": 90, "low": 180}
    sla_within:  dict[str, int] = {"critical": 0, "high": 0, "medium": 0, "low": 0}
    sla_breached: dict[str, int] = {"critical": 0, "high": 0, "medium": 0, "low": 0}

    # ── Secret Detection Trend ─────────────────────────────────
    # Per-target, per-scan: count of TruffleHog findings over time (last 60 scans)
    secret_trend: list[dict] = []
    secret_trend_by_target: dict[str, list[dict]] = {}

    # ── Weighted Risk Score Trend ──────────────────────────────
    # Score = critical*10 + high*7 + medium*4 + low*1, per scan
    RISK_WEIGHTS = {"critical": 10, "high": 7, "medium": 4, "low": 1}
    risk_trend: list[dict] = []

    # ── Suppression Rate ───────────────────────────────────────
    total_suppressed = 0
    suppressed_by_sev: dict[str, int] = {"critical": 0, "high": 0, "medium": 0, "low": 0}
    suppressed_by_tool: dict[str, int] = {}
    suppressed_by_target: dict[str, list[dict]] = {}

    # ── Recurrence + First-Time Fix Rate ──────────────────────
    # recurrence: fingerprints that were resolved but reappeared
    # first_time_fix: fingerprints resolved within 1 scan cycle of first appearance
    recurrence_count   = 0
    first_time_fixed   = 0
    total_resolved     = 0

    for _target, scan_list in by_target.items():
        eligible = [
            (s, d) for s, d in scan_list
            if s.get("scan_type") not in ("stig",)
        ]
        if not eligible:
            continue
        eligible_sorted = sorted(eligible, key=lambda x: x[0].get("timestamp", ""))

        # Secret trend for this target
        t_secret_trend: list[dict] = []
        for scan_meta, scan_dir in eligible_sorted:
            ts = (scan_meta.get("timestamp") or "")[:10]
            if not ts:
                continue
            # Risk score uses only scan_meta — compute before attempting findings parse
            s = scan_meta
            risk_trend.append({
                "date":     ts,
                "target":   _target,
                "score":    (s.get("critical", 0) * RISK_WEIGHTS["critical"]
                             + s.get("high",     0) * RISK_WEIGHTS["high"]
                             + s.get("medium",   0) * RISK_WEIGHTS["medium"]
                             + s.get("low",      0) * RISK_WEIGHTS["low"]),
                "critical": s.get("critical", 0),
                "high":     s.get("high",     0),
                "medium":   s.get("medium",   0),
                "low":      s.get("low",      0),
            })
            try:
                sf = parsers.parse_scan_findings(scan_dir)
            except Exception:
                continue
            secrets = [
                f for sev in ("critical", "high", "medium", "low")
                for f in sf.get(f"{sev}_findings", [])
                if f.get("tool") == "TruffleHog"
            ]
            verified   = sum(1 for f in secrets if f.get("severity") == "critical")
            unverified = len(secrets) - verified
            t_secret_trend.append({
                "date": ts, "target": _target,
                "verified": verified, "unverified": unverified,
                "total": len(secrets),
            })
        secret_trend.extend(t_secret_trend)
        if t_secret_trend:
            secret_trend_by_target[_target] = t_secret_trend

        # Suppression: union across all scans for target (latest scan may have empty file)
        seen_sup_keys: set[tuple] = set()
        target_suppressions: list[dict] = []
        for _, scan_dir in reversed(eligible_sorted):  # newest → oldest
            try:
                batch = parsers.parse_suppressed_findings(scan_dir)
            except Exception:
                continue
            for sup in batch:
                key = (sup.get("type", ""), sup.get("value", ""))
                if key not in seen_sup_keys:
                    seen_sup_keys.add(key)
                    target_suppressions.append(sup)
        total_suppressed += len(target_suppressions)
        if target_suppressions:
            suppressed_by_target[_target] = [
                {
                    "value":       sup.get("value", ""),
                    "tool":        sup.get("tool", "Unknown"),
                    "type":        sup.get("type", ""),
                    "severity":    sup.get("severity", ""),
                    "reason":      sup.get("reason", ""),
                    "approved_by": sup.get("approved_by", ""),
                }
                for sup in target_suppressions
            ]
        for sup in target_suppressions:
            raw_sev = (sup.get("severity") or "").lower()
            if raw_sev in suppressed_by_sev:
                suppressed_by_sev[raw_sev] += 1
            tool_key = (sup.get("tool") or "Unknown")
            suppressed_by_tool[tool_key] = suppressed_by_tool.get(tool_key, 0) + 1

        if len(eligible_sorted) < 2:
            continue

        # Build per-scan fingerprint sets for recurrence + SLA + first-time-fix
        fp_first_seen: dict[str, tuple[str, str]] = {}  # fp -> (timestamp, sev)
        fp_resolved_at: dict[str, str] = {}             # fp -> resolved_scan_ts
        fp_ever_resolved: set[str]     = set()
        # Pending-resolution: fp disappeared last scan but we require one more
        # consecutive absence before counting it as a genuine fix (filters scanner noise).
        fp_pending_resolution: dict[str, str] = {}      # fp -> scan_ts when it first disappeared
        prev_fps: set[str]             = set()
        prev_ts_str: str               = ""

        for scan_meta, scan_dir in eligible_sorted:
            ts = scan_meta.get("timestamp", "")
            if not ts:
                continue
            try:
                sf = parsers.parse_scan_findings(scan_dir)
            except Exception:
                prev_fps = set()
                prev_ts_str = ts
                fp_pending_resolution.clear()
                continue
            curr_fps: set[str] = set()
            curr_fp_sev: dict[str, str] = {}
            for sev in ("critical", "high", "medium", "low"):
                for f in sf.get(f"{sev}_findings", []):
                    fp = f"{f.get('tool', '')}::{f.get('id', '')}::{f.get('package', '')}"
                    if fp:
                        curr_fps.add(fp)
                        curr_fp_sev[fp] = sev
                        if fp not in fp_first_seen:
                            fp_first_seen[fp] = (ts, sev)
                        # Recurrence: was resolved before, now back
                        if fp in fp_ever_resolved:
                            recurrence_count += 1
                            fp_ever_resolved.discard(fp)
                        # Reappeared before confirmation — cancel the pending resolution
                        fp_pending_resolution.pop(fp, None)

            # ── Two-scan confirmation: a finding must be absent from 2 consecutive
            # scans before it counts as resolved. Single-scan absences (scanner noise,
            # transient fetch failures, DB updates) are ignored. ──────────────────
            if prev_fps:
                disappeared_now = prev_fps - curr_fps

                # Confirm resolutions that were pending from the previous scan
                # (absent LAST scan AND still absent NOW → confirmed resolved)
                sla_seen_this_transition: set[str] = set()
                for fp, absent_since_ts in list(fp_pending_resolution.items()):
                    if fp not in curr_fps:  # still absent — confirmed
                        del fp_pending_resolution[fp]
                        if fp in fp_first_seen:
                            first_ts, sev = fp_first_seen[fp]
                            fp_ever_resolved.add(fp)
                            fp_resolved_at[fp] = absent_since_ts
                            total_resolved += 1
                            if prev_ts_str == first_ts:
                                first_time_fixed += 1
                            # SLA — deduplicate by CVE-ID::package (strip tool prefix)
                            fp_parts = fp.split('::', 2)
                            sla_key = f"{fp_parts[1]}::{fp_parts[2]}" if len(fp_parts) == 3 and fp_parts[1] else fp
                            if sla_key not in sla_seen_this_transition:
                                sla_seen_this_transition.add(sla_key)
                                try:
                                    first_dt = datetime.fromisoformat(first_ts.replace("Z", "+00:00"))
                                    res_dt   = datetime.fromisoformat(absent_since_ts.replace("Z", "+00:00"))
                                    if first_dt.tzinfo is None:
                                        first_dt = first_dt.replace(tzinfo=timezone.utc)
                                    if res_dt.tzinfo is None:
                                        res_dt = res_dt.replace(tzinfo=timezone.utc)
                                    days = (res_dt - first_dt).total_seconds() / 86400
                                    threshold = SLA_DAYS.get(sev, 180)
                                    if days <= threshold:
                                        sla_within[sev]   += 1
                                    else:
                                        sla_breached[sev] += 1
                                except ValueError:
                                    pass
                    # else: reappeared → already removed from pending_resolution above

                # Stage newly disappeared fingerprints as pending (not yet confirmed)
                for fp in disappeared_now:
                    if fp not in fp_ever_resolved and fp in fp_first_seen:
                        fp_pending_resolution[fp] = ts

            prev_fps    = curr_fps
            prev_ts_str = ts

    # Aggregate SLA
    sla_by_sev = {}
    for sev in ("critical", "high", "medium", "low"):
        w = sla_within[sev]
        b = sla_breached[sev]
        total_sla = w + b
        sla_by_sev[sev] = {
            "within": w,
            "breached": b,
            "pct": round(w / total_sla * 100, 1) if total_sla > 0 else None,
            "sla_days": SLA_DAYS[sev],
        }
    sla_total_within  = sum(sla_within.values())
    sla_total_breached = sum(sla_breached.values())
    sla_total = sla_total_within + sla_total_breached
    sla_overall_pct = round(sla_total_within / sla_total * 100, 1) if sla_total > 0 else None

    # Aggregate risk trend (last 60 data points per target, sorted)
    by_target_trend: dict[str, list[dict]] = {}
    for pt in risk_trend:
        by_target_trend.setdefault(pt["target"], []).append(pt)

    risk_trend_sorted: list[dict] = []
    for target, pts in by_target_trend.items():
        pts_sorted = sorted(pts, key=lambda x: x["date"])[-60:]
        risk_trend_sorted.extend(pts_sorted)

    risk_trend_sorted.sort(key=lambda x: (x["target"], x["date"]))
    # Aggregate secret trend (last 60 total entries)
    secret_trend_sorted = sorted(secret_trend, key=lambda x: x["date"])[-60:]

    # Suppression rate vs fixed
    total_fixed_or_suppressed = total_resolved + total_suppressed
    suppression_rate_pct = (
        round(total_suppressed / total_fixed_or_suppressed * 100, 1)
        if total_fixed_or_suppressed > 0 else None
    )

    # Recurrence rate
    recurrence_rate_pct = (
        round(recurrence_count / total_resolved * 100, 1)
        if total_resolved > 0 else None
    )

    # First-time fix rate
    first_time_fix_pct = (
        round(first_time_fixed / total_resolved * 100, 1)
        if total_resolved > 0 else None
    )

    # ── MTTR (Mean Time to Remediate) + MTTD (Mean Time to Detect) ──────────────    # by_target lists are already sorted newest-first from the loop above
    all_remediation_days: list[float] = []
    all_remediation_days_by_sev: dict[str, list[float]] = {"critical": [], "high": [], "medium": [], "low": []}
    all_detection_days: list[float] = []
    all_detection_days_by_sev: dict[str, list[float]] = {"critical": [], "high": [], "medium": [], "low": []}
    target_mttr: dict[str, float] = {}
    # Per-target detailed metrics for app-level filtering
    metrics_by_target: dict[str, dict] = {}
    for _target, scan_list in by_target.items():
        # Only full and nightly scans carry consistent vulnerability data
        eligible = [
            (s, d) for s, d in scan_list
            if s.get("scan_type") not in ("stig",)
        ]
        if len(eligible) < 2:
            continue
        # eligible is newest-first; reverse to chronological order
        scans_to_check = list(reversed(eligible))  # oldest → newest

        # Build fingerprint -> first_seen_ts + severity mapping
        first_seen: dict[str, str] = {}
        first_seen_sev: dict[str, str] = {}
        known_fps: set[str] = set()
        prev_ts: str | None = None
        t_detection_days: list[float] = []
        t_detection_days_by_sev: dict[str, list[float]] = {"critical": [], "high": [], "medium": [], "low": []}
        for scan_meta, scan_dir in scans_to_check:
            ts = scan_meta.get("timestamp", "")
            if not ts:
                continue
            try:
                scan_findings = parsers.parse_scan_findings(scan_dir)
            except Exception:
                prev_ts = ts
                continue
            current_fps: set[str] = set()
            current_fp_sev: dict[str, str] = {}
            for sev in ("critical", "high", "medium", "low"):
                for f in scan_findings.get(f"{sev}_findings", []):
                    fp = f"{f.get('tool', '')}::{f.get('id', '')}::{f.get('package', '')}"
                    if fp:
                        if fp not in first_seen:
                            first_seen[fp] = ts
                            first_seen_sev[fp] = sev
                        current_fps.add(fp)
                        current_fp_sev[fp] = sev
            # MTTD: record detection window for newly-appearing findings
            if prev_ts:
                try:
                    prev_dt = datetime.fromisoformat(prev_ts.replace("Z", "+00:00"))
                    curr_dt = datetime.fromisoformat(ts.replace("Z", "+00:00"))
                    if prev_dt.tzinfo is None:
                        prev_dt = prev_dt.replace(tzinfo=timezone.utc)
                    if curr_dt.tzinfo is None:
                        curr_dt = curr_dt.replace(tzinfo=timezone.utc)
                    detect_days = (curr_dt - prev_dt).total_seconds() / 86400
                    if detect_days >= 0:
                        for fp in (current_fps - known_fps):
                            sev = current_fp_sev.get(fp, "")
                            all_detection_days.append(detect_days)
                            t_detection_days.append(detect_days)
                            if sev in all_detection_days_by_sev:
                                all_detection_days_by_sev[sev].append(detect_days)
                            if sev in t_detection_days_by_sev:
                                t_detection_days_by_sev[sev].append(detect_days)
                except ValueError:
                    pass
            known_fps |= current_fps
            prev_ts = ts

        # Latest scan fingerprint set (scans_to_check[-1] is the newest)
        latest_meta, latest_dir = scans_to_check[-1]
        latest_ts = latest_meta.get("timestamp", "")
        if not latest_ts:
            continue
        latest_set: set[str] = set()
        try:
            latest_findings = parsers.parse_scan_findings(latest_dir)
        except Exception:
            continue
        for sev in ("critical", "high", "medium", "low"):
            for f in latest_findings.get(f"{sev}_findings", []):
                fp = f"{f.get('tool', '')}::{f.get('id', '')}::{f.get('package', '')}"
                if fp:
                    latest_set.add(fp)

        try:
            latest_dt = datetime.fromisoformat(latest_ts.replace("Z", "+00:00"))
            if latest_dt.tzinfo is None:
                latest_dt = latest_dt.replace(tzinfo=timezone.utc)
        except ValueError:
            continue

        target_days: list[float] = []
        t_remediation_days_by_sev: dict[str, list[float]] = {"critical": [], "high": [], "medium": [], "low": []}
        for fp, first_ts in first_seen.items():
            if fp not in latest_set:
                try:
                    first_dt = datetime.fromisoformat(first_ts.replace("Z", "+00:00"))
                    if first_dt.tzinfo is None:
                        first_dt = first_dt.replace(tzinfo=timezone.utc)
                    days = (latest_dt - first_dt).total_seconds() / 86400
                    if days >= 0:
                        all_remediation_days.append(days)
                        target_days.append(days)
                        sev = first_seen_sev.get(fp, "")
                        if sev in all_remediation_days_by_sev:
                            all_remediation_days_by_sev[sev].append(days)
                        if sev in t_remediation_days_by_sev:
                            t_remediation_days_by_sev[sev].append(days)
                except ValueError:
                    pass
        if target_days:
            target_mttr[_target] = round(sum(target_days) / len(target_days), 1)

        t_mttr = round(sum(target_days) / len(target_days), 1) if target_days else None
        t_mttd = round(sum(t_detection_days) / len(t_detection_days), 1) if t_detection_days else None
        metrics_by_target[_target] = {
            "mttr_days": t_mttr,
            "mttr_by_severity": {
                s: round(sum(v) / len(v), 1) if v else None
                for s, v in t_remediation_days_by_sev.items()
            },
            "mttd_days": t_mttd,
            "mttd_by_severity": {
                s: round(sum(v) / len(v), 1) if v else None
                for s, v in t_detection_days_by_sev.items()
            },
        }

    mttr_days = (
        round(sum(all_remediation_days) / len(all_remediation_days), 1)
        if all_remediation_days else None
    )
    mttr_by_severity = {
        sev: round(sum(days) / len(days), 1) if days else None
        for sev, days in all_remediation_days_by_sev.items()
    }
    mttd_days = (
        round(sum(all_detection_days) / len(all_detection_days), 1)
        if all_detection_days else None
    )
    mttd_by_severity = {
        sev: round(sum(days) / len(days), 1) if days else None
        for sev, days in all_detection_days_by_sev.items()
    }
    fastest_remediator = (
        min(target_mttr.items(), key=lambda x: x[1])
        if target_mttr else None
    )

    result = {
        "trend": trend,
        "by_tool": sorted(
            [
                {
                    "tool":     k,
                    "critical": v["critical"],
                    "high":     v["high"],
                    "medium":   v["medium"],
                    "low":      v["low"],
                    "total":    v["total"],
                    "top_app":  max(tool_app_counts.get(k, {"":0}).items(), key=lambda x: x[1])[0],
                }
                for k, v in tool_counts.items()
            ],
            key=lambda x: -x["total"],
        ),
        "fix_rate": {"with_fix": total_with_fix, "without_fix": total_without_fix},
        "mttr_days": mttr_days,
        "mttr_by_severity": mttr_by_severity,
        "mttd_days": mttd_days,
        "mttd_by_severity": mttd_by_severity,
        "metrics_by_target": metrics_by_target,
        "fastest_remediator": {"target": fastest_remediator[0], "mttr_days": fastest_remediator[1]} if fastest_remediator else None,
        "top_cves": [
            {
                "cve_id":   k,
                "count":    v["count"],
                "severity": v["severity"],
                "title":    v["title"],
                "apps":     sorted(v["apps"]),
                "sources":  sorted(v["sources"]),
            }
            for k, v in top_cves
        ],
        "scan_frequency": scan_frequency,
        "merges_to_main": merges_to_main,
        "total_merges_to_main": total_merges_to_main,
        "metrics_filtered": metrics_filtered,
        "monitored_count":  monitored_count,
        # ── New metrics ──────────────────────────────────────
        "sla_compliance": {
            "overall_pct":     sla_overall_pct,
            "total_within":    sla_total_within,
            "total_breached":  sla_total_breached,
            "by_severity":     sla_by_sev,
        },
        "risk_trend":            risk_trend_sorted,
        "secret_trend":          secret_trend_sorted,
        "secret_trend_by_target": secret_trend_by_target,
        "suppression": {
            "total":       total_suppressed,
            "by_severity": suppressed_by_sev,
            "by_tool":     suppressed_by_tool,
            "by_target":   suppressed_by_target,
            "rate_pct":    suppression_rate_pct,
        },
        "accepted_risk_pct": (
            round(total_suppressed_instances /
                  (total_with_fix + total_without_fix + total_suppressed_instances) * 100, 1)
            if (total_with_fix + total_without_fix + total_suppressed_instances) > 0 else None
        ),
        "recurrence_rate_pct":    recurrence_rate_pct,
        "first_time_fix_pct":     first_time_fix_pct,
        "total_resolved":         total_resolved,
    }

    _metrics_cache    = result
    _metrics_cache_ts = now
    return result


# ── Settings ──────────────────────────────────────────────────

@app.get("/api/settings/approved-images")
def approved_images(response: Response):
    _sec_headers(response)
    try:
        content = APPROVED_IMAGES_FILE.read_text(encoding="utf-8")
    except OSError:
        content = ""
    return {"content": content}


# ── GitHub Metrics ────────────────────────────────────────────

@app.get("/api/github-metrics")
async def get_github_metrics(response: Response):
    _sec_headers(response)
    global _gh_metrics_cache, _gh_metrics_cache_ts
    now = time.monotonic()
    if _gh_metrics_cache and (now - _gh_metrics_cache_ts) < _GH_METRICS_TTL:
        return _gh_metrics_cache

    cfg   = _read_github_config()
    token = cfg.get("token", "")
    repos = cfg.get("repos") or []

    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    all_scans = [parsers.load_scan(d, EPYON_ROOT) for d in scan_dirs]
    monitored_set = set(_load_monitored_apps()) or None

    result = await github_metrics.fetch_all(token, repos, all_scans, monitored_set)

    _gh_metrics_cache    = result
    _gh_metrics_cache_ts = now
    _save_gh_snapshot(result)
    return result


@app.get("/api/github-signals-history")
def get_github_signals_history(response: Response):
    _sec_headers(response)
    return _load_gh_history()


# ── GitHub config ─────────────────────────────────────────────

@app.get("/api/github/config")
def github_config_get(response: Response):
    _sec_headers(response)
    cfg = _read_github_config()
    token = cfg.get("token", "")
    masked = re.sub(r"(?<=.{7}).(?=.{4})", "*", token) if token else ""

    # Mask extra tokens too — return repo list + masked hint per entry
    extra = []
    for entry in (cfg.get("extra_tokens") or []):
        t = entry.get("token", "")
        extra.append({
            "repos":      entry.get("repos") or [],
            "token_set":  bool(t),
            "token_hint": re.sub(r"(?<=.{7}).(?=.{4})", "*", t) if t else "",
        })

    return {
        "token_set":    bool(token),
        "token_hint":   masked,
        "repos":        cfg.get("repos") or [],
        "extra_tokens": extra,
        "last_sync":    cfg.get("last_sync"),
    }


@app.post("/api/github/config")
async def github_config_post(request: Request, response: Response):
    _sec_headers(response)
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "Invalid JSON")

    cfg = _read_github_config()
    new_token = (body.get("token") or "").strip()
    if new_token and new_token != "KEEP_EXISTING":
        if not _TOKEN_RE.match(new_token):
            raise HTTPException(400, "Token does not look like a valid GitHub token")
        cfg["token"] = new_token

    if isinstance(body.get("repos"), list):
        cfg["repos"] = [
            r.strip() for r in body["repos"]
            if isinstance(r, str) and _REPO_RE.match(r.strip())
        ]

    # extra_tokens: list of {"repos": [...], "token": "ghp_..."|"KEEP_EXISTING"}
    if isinstance(body.get("extra_tokens"), list):
        existing_extras = {i: e for i, e in enumerate(cfg.get("extra_tokens") or [])}
        new_extras = []
        for idx, entry in enumerate(body["extra_tokens"]):
            t = (entry.get("token") or "").strip()
            repos_list = [
                r.strip() for r in (entry.get("repos") or [])
                if isinstance(r, str) and _REPO_RE.match(r.strip())
            ]
            if t == "KEEP_EXISTING":
                # Preserve previously stored token for this slot
                t = (existing_extras.get(idx) or {}).get("token", "")
            elif t and not _TOKEN_RE.match(t):
                raise HTTPException(400, f"Extra token at index {idx} is not a valid GitHub token")
            if repos_list or t:
                new_extras.append({"repos": repos_list, "token": t})
        cfg["extra_tokens"] = new_extras

    _write_github_config(cfg)
    return {"ok": True}


# ── GitHub sync ───────────────────────────────────────────────

@app.post("/api/github/sync")
async def github_sync_post(response: Response):
    _sec_headers(response)
    state = github_sync.get_sync_state()
    if state["status"] == "running":
        return {"ok": False, "message": "Sync already in progress"}
    await github_sync.trigger_sync(
        GITHUB_CONFIG_FILE, EPYON_ROOT, parsers.find_scan_dirs,
        on_complete=_invalidate_scan_cache,
    )
    return {"ok": True, "message": "Sync started"}


# ── STIG History / MTTR ──────────────────────────────────────

@app.get("/api/stig/history")
def stig_history(
    response: Response,
    app: Optional[str] = Query(None),
    slug: Optional[str] = Query(None),
):
    """Return STIG control status history across all scans for an app, with MTTR."""
    _sec_headers(response)
    scan_dirs = _cached_find_scan_dirs()

    # Group scan dirs by app name, extract date/time from scan_id
    app_scans: dict[str, list[dict]] = {}
    for d in scan_dirs:
        m = _APP_SCAN_RE.match(d.name)
        if not m:
            continue
        app_name, date_str, time_str = m.group(1), m.group(2), m.group(3)
        app_scans.setdefault(app_name, []).append(
            {"scan_id": d.name, "date": date_str, "time": time_str, "path": d}
        )

    # Apps that have at least one stig-results-*.json file
    stig_apps = sorted(
        a for a, scans in app_scans.items()
        if any(list(s["path"].glob("stig-results-*.json")) for s in scans)
    )

    _empty_summary = {
        "total_controls": 0, "ever_open": 0,
        "remediated": 0, "avg_mttr_days": None, "currently_open": 0,
    }
    _base = {"apps": stig_apps, "app": None, "slugs": [], "slug": slug,
             "scans": [], "controls": [], "summary": _empty_summary}

    selected_app = app or (stig_apps[0] if stig_apps else None)
    if not selected_app:
        return _base

    # All slugs found for this app
    all_slugs: set[str] = set()
    for s in app_scans.get(selected_app, []):
        for f in s["path"].glob("stig-results-*.json"):
            all_slugs.add(f.stem[len("stig-results-"):])

    selected_slugs = {slug} if (slug and slug in all_slugs) else all_slugs

    # Scans for this app that have STIG files, in chronological order
    stig_scans = sorted(
        [
            s for s in app_scans.get(selected_app, [])
            if any(
                (s["path"] / f"stig-results-{sl}.json").exists()
                for sl in selected_slugs
            )
        ],
        key=lambda s: (s["date"], s["time"]),
    )

    # Deduplicate: keep only the latest STIG scan per calendar date
    _latest_per_date: dict[str, dict] = {}
    for s in stig_scans:
        _latest_per_date[s["date"]] = s
    stig_scans = list(_latest_per_date.values())

    if not stig_scans:
        return {**_base, "apps": stig_apps, "app": selected_app, "slugs": sorted(all_slugs)}

    # Build control metadata (title/severity) and the canonical vuln_id set per slug
    # (using the most-recent controls file per slug so the benchmark version is consistent)
    control_meta: dict[str, dict] = {}
    slug_control_ids: dict[str, set[str]] = {}
    for s in reversed(stig_scans):
        for sl in selected_slugs:
            cf = s["path"] / f"stig-controls-{sl}.json"
            if not cf.exists():
                continue
            try:
                cd = json.loads(cf.read_text(encoding="utf-8"))
                ids_for_slug: set[str] = set()
                for c in cd.get("controls", []):
                    vid = c.get("vuln_id", "")
                    if not vid:
                        continue
                    if vid not in control_meta:
                        control_meta[vid] = {
                            "title":     c.get("title", ""),
                            "severity":  c.get("severity", ""),
                            "stig_name": cd.get("stig_name", sl),
                        }
                    ids_for_slug.add(vid)
                # Only use the first (most-recent) controls file per slug
                if sl not in slug_control_ids:
                    slug_control_ids[sl] = ids_for_slug
            except Exception:
                pass

    # Build per-vuln_id timeline, filling gaps with "Not Reviewed" for any
    # benchmark control that a scan ran but didn't explicitly produce a result for.
    # This guarantees every scan column is complete and the matrix is consistent.
    seen_entries: set[tuple] = set()
    vuln_timelines: dict[str, list[dict]] = {}
    for s in stig_scans:
        for sl in selected_slugs:
            rf = s["path"] / f"stig-results-{sl}.json"
            if not rf.exists():
                continue
            try:
                raw = json.loads(rf.read_text(encoding="utf-8"))
                results = raw.get("assessments", raw) if "assessments" in raw else raw
            except Exception:
                continue
            # Explicit assessment results
            for vid, data in results.items():
                key = (s["scan_id"], vid)
                if key not in seen_entries:
                    seen_entries.add(key)
                    vuln_timelines.setdefault(vid, []).append({
                        "scan_id": s["scan_id"],
                        "date":    s["date"],
                        "status":  data.get("status", "Not Reviewed"),
                        "evidence": data.get("evidence", ""),
                        "confidence": data.get("confidence", None),
                    })
            # Fill any benchmark controls absent from this scan's results
            for vid in slug_control_ids.get(sl, set()):
                key = (s["scan_id"], vid)
                if key not in seen_entries:
                    seen_entries.add(key)
                    vuln_timelines.setdefault(vid, []).append({
                        "scan_id": s["scan_id"],
                        "date":    s["date"],
                        "status":  "Not Reviewed",
                        "evidence": "Control not assessed in this scan.",
                        "confidence": None,
                    })

    # Compute MTTR and assemble output
    controls_out: list[dict] = []
    total_mttr: list[int] = []
    currently_open = ever_open = remediated = currently_satisfied = 0

    for vid, timeline in sorted(vuln_timelines.items()):
        tl = sorted(timeline, key=lambda t: t["date"])
        first_open = None
        first_closed = None  # first "Not a Finding" AFTER first_open
        for entry in tl:
            if entry["status"] == "Open" and first_open is None:
                first_open = entry["date"]
            # Only count as remediated if it closed AFTER it was first opened
            if (entry["status"] == "Not a Finding"
                    and first_open is not None
                    and first_closed is None
                    and entry["date"] >= first_open):
                first_closed = entry["date"]

        mttr_days = None
        if first_open:
            ever_open += 1
        if first_open and first_closed:
            remediated += 1
            d1 = datetime.fromisoformat(first_open).date()
            d2 = datetime.fromisoformat(first_closed).date()
            mttr_days = (d2 - d1).days
            total_mttr.append(mttr_days)

        latest_status = tl[-1]["status"] if tl else "Not Reviewed"
        if latest_status == "Open":
            currently_open += 1
        if latest_status in ("Not a Finding", "Not Applicable"):
            currently_satisfied += 1

        meta = control_meta.get(vid, {})
        controls_out.append({
            "vuln_id":       vid,
            "title":         meta.get("title", ""),
            "severity":      meta.get("severity", ""),
            "stig_name":     meta.get("stig_name", ""),
            "timeline":      tl,
            "first_open":    first_open,
            "first_closed":  first_closed,
            "mttr_days":     mttr_days,
            "latest_status": latest_status,
        })

    avg_mttr = round(sum(total_mttr) / len(total_mttr), 1) if total_mttr else None
    scans_out = [
        {"scan_id": s["scan_id"], "date": s["date"], "label": s["date"][5:]}
        for s in stig_scans
    ]

    return {
        "apps":     stig_apps,
        "app":      selected_app,
        "slugs":    sorted(all_slugs),
        "slug":     slug,
        "scans":    scans_out,
        "controls": controls_out,
        "summary": {
            "total_controls":     len(controls_out),
            "ever_open":          ever_open,
            "remediated":         remediated,
            "avg_mttr_days":      avg_mttr,
            "currently_open":     currently_open,
            "currently_satisfied": currently_satisfied,
        },
    }


@app.get("/api/github/sync")
def github_sync_status(response: Response):
    _sec_headers(response)
    return github_sync.get_sync_state()


# ── AI / OpenAI config ────────────────────────────────────────

@app.get("/api/ai/config")
def ai_config_get(response: Response):
    _sec_headers(response)
    cfg = openai_summary.read_ai_config()
    key = cfg.get("api_key", "")
    masked = re.sub(r"(?<=.{7}).(?=.{4})", "*", key) if key else ""
    return {
        "key_set":  bool(key),
        "key_hint": masked,
        # `model` reflects only the UI-managed value so the Settings dropdown
        # stays in sync with what was actually saved. Returning the resolved
        # effective model here would surface a non-allowlisted env value (e.g.
        # gemma4:26b) that the dropdown can't represent; the browser would then
        # fall back to its first <option> and a subsequent Save would persist
        # that, clobbering the env-configured model. `active_model` exposes the
        # resolved model (UI > OPENAI_MODEL > default) for display only.
        "model":        cfg.get("model") or openai_summary.DEFAULT_MODEL,
        "active_model": openai_summary.get_model(),
        # The configured custom endpoint (UI value), plus the effective base URL
        # actually in use (UI > OPENAI_BASE_URL env > OpenAI default) for
        # display. Empty string means the public OpenAI API.
        "base_url":        cfg.get("base_url", ""),
        "active_base_url": openai_summary.get_base_url() or "",
    }


@app.post("/api/ai/config")
async def ai_config_post(request: Request, response: Response):
    _sec_headers(response)
    _audit(request, "ai_config_changed")
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "Invalid JSON")

    cfg = openai_summary.read_ai_config()
    new_key = (body.get("api_key") or "").strip()
    if new_key and new_key != "KEEP_EXISTING":
        # Accept real OpenAI secret keys (sk-...) as well as arbitrary tokens
        # used by self-hosted, OpenAI-compatible gateways (vLLM, LiteLLM, ...).
        # Reject only obviously unsafe values (whitespace / shell metacharacters
        # / control chars) rather than locking to the OpenAI key format.
        if len(new_key) > 200 or re.search(r"[\s;&|`$()<>'\"\\]", new_key):
            raise HTTPException(400, "api_key contains invalid characters")
        cfg["api_key"] = new_key

    if "base_url" in body:
        base_url = (body.get("base_url") or "").strip()
        if base_url:
            # Allow https:// for public providers and http:// for in-cluster /
            # self-hosted endpoints (e.g. http://ollama.ollama:11434/v1).
            if not re.match(r"^https?://[A-Za-z0-9.\-]+(:\d+)?(/[\w./\-]*)?$", base_url):
                raise HTTPException(400, "base_url must be a valid http(s) URL")
            # SSRF / credential-exfil guard: only the public OpenAI API,
            # cluster-internal/private hosts, or an operator allowlist
            # (EPYON_AI_ALLOWED_HOSTS) may be targeted, so this endpoint cannot
            # be used to redirect a stored API key to an attacker host.
            if not openai_summary.base_url_allowed(base_url):
                raise HTTPException(
                    400,
                    "base_url host is not allowed; use the public OpenAI API, a "
                    "cluster-internal endpoint, or add it to EPYON_AI_ALLOWED_HOSTS",
                )
            cfg["base_url"] = base_url
        else:
            cfg.pop("base_url", None)  # explicit clear → fall back to env/default

    if body.get("model"):
        model = str(body["model"]).strip()
        # Accept curated OpenAI models and any self-hosted model id
        # (e.g. gemma4:26b, llama3.1:8b). Validate the character set so the
        # value is safe to use in request bodies and file paths.
        if not re.match(r"^[A-Za-z0-9._:\-/]{1,100}$", model):
            raise HTTPException(400, "model contains invalid characters")
        cfg["model"] = model

    openai_summary.write_ai_config(cfg)
    return {"ok": True}


# ── NVD API config ────────────────────────────────────────────

NVD_CONFIG_FILE = Path("data/nvd-config.json")

def read_nvd_config() -> dict:
    """Read NVD API key config from data/nvd-config.json."""
    if NVD_CONFIG_FILE.exists():
        try:
            with open(NVD_CONFIG_FILE) as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def write_nvd_config(cfg: dict):
    """Write NVD API key config to data/nvd-config.json."""
    NVD_CONFIG_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(NVD_CONFIG_FILE, "w") as f:
        json.dump(cfg, f, indent=2)

@app.get("/api/nvd/config")
def nvd_config_get(response: Response):
    _sec_headers(response)
    # Check env var first (takes priority)
    env_key = os.getenv("NVD_API_KEY", "")
    if env_key:
        masked = re.sub(r"(?<=.{4}).(?=.{4})", "*", env_key) if env_key else ""
        return {
            "key_set": True,
            "key_hint": masked,
            "from_env": True,
        }
    
    # Otherwise check saved config
    cfg = read_nvd_config()
    key = cfg.get("api_key", "")
    masked = re.sub(r"(?<=.{4}).(?=.{4})", "*", key) if key else ""
    return {
        "key_set": bool(key),
        "key_hint": masked,
        "from_env": False,
    }

@app.post("/api/nvd/config")
async def nvd_config_post(request: Request, response: Response):
    _sec_headers(response)
    _audit(request, "nvd_config_changed")
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "Invalid JSON")

    cfg = read_nvd_config()
    new_key = (body.get("api_key") or "").strip()
    if new_key and new_key != "KEEP_EXISTING":
        # NVD API keys are UUIDs with dashes
        if not re.match(r"^[a-f0-9]{8}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{12}$", new_key, re.IGNORECASE):
            raise HTTPException(400, "api_key must be a valid NVD API key (UUID format)")
        cfg["api_key"] = new_key
    
    write_nvd_config(cfg)
    return {"ok": True}


# ── Executive summary ─────────────────────────────────────────

@app.post("/api/scans/{scan_id}/executive-summary")
async def exec_summary(scan_id: str, response: Response):
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")

    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")

    scan_meta = parsers.load_scan(matched, EPYON_ROOT)
    findings  = parsers.parse_scan_findings(matched)

    try:
        summary = await openai_summary.generate_summary(scan_id, scan_meta, findings)
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))
    except Exception as exc:
        raise HTTPException(502, f"OpenAI request failed: {exc}")

    return {"scan_id": scan_id, "summary": summary}


@app.post("/api/scans/{scan_id}/technical-summary")
async def technical_summary(scan_id: str, response: Response):
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")

    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")

    scan_meta = parsers.load_scan(matched, EPYON_ROOT)
    findings  = parsers.parse_scan_findings(matched)

    try:
        summary = await openai_summary.generate_technical_summary(scan_id, scan_meta, findings)
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))
    except Exception as exc:
        raise HTTPException(502, f"OpenAI request failed: {exc}")

    return {"scan_id": scan_id, "summary": summary}


@app.post("/api/scans/{scan_id}/scorecard")
async def calculate_scorecard(scan_id: str, response: Response):
    """Calculate security scorecard for an existing scan"""
    _sec_headers(response)
    if not _SAFE_ID_RE.match(scan_id):
        raise HTTPException(400, "Invalid scan_id")

    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    matched = next((d for d in scan_dirs if d.name == scan_id), None)
    if not matched:
        raise HTTPException(404, "Scan not found")

    scan_dir = EPYON_ROOT / "scans" / scan_id
    trl_script = SCRIPTS_DIR / "generate-trl-score.py"

    if not trl_script.exists():
        raise HTTPException(500, "TRL scoring script not found")

    # Run TRL calculation
    import subprocess
    result = subprocess.run(
        ["python3", str(trl_script), "--scan-dir", str(scan_dir)],
        capture_output=True,
        text=True,
        timeout=60,
    )

    if result.returncode != 0:
        raise HTTPException(500, f"Score calculation failed: {result.stderr}")

    # Read and return the result
    trl_file = scan_dir / "trl-assessment.json"
    if not trl_file.exists():
        raise HTTPException(500, "Score assessment file not generated")

    try:
        trl_data = json.loads(trl_file.read_text(encoding="utf-8"))
    except Exception as exc:
        raise HTTPException(500, f"Failed to parse score results: {exc}")

    return {"scan_id": scan_id, **trl_data}


@app.post("/api/executive-summary")
async def global_exec_summary(response: Response):
    _sec_headers(response)
    hidden    = _load_hidden_apps()
    monitored = _load_monitored_apps()
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)

    # Collect all scan dirs per target, then choose the latest comprehensive scan
    # (full/nightly) for vulnerability data. Fall back to absolute latest if no
    # comprehensive scan exists for a target.
    by_target_all: dict[str, list[tuple]] = {}  # target -> [(dir, timestamp), ...]
    for d in scan_dirs:
        meta = parsers.parse_dir_name(d.name)
        target = meta["target"]
        if target in hidden:
            continue
        by_target_all.setdefault(target, []).append((d, meta["timestamp"]))

    if not by_target_all:
        raise HTTPException(404, "No scans found")

    def _best_scan_dir(entries: list[tuple]) -> object:
        entries_sorted = sorted(entries, key=lambda x: x[1], reverse=True)
        for d, _ in entries_sorted:
            st = (parsers._read_json(d / "scan-metadata.json") or {}).get("scan_type", "")
            if st in _COMPREHENSIVE_SCAN_TYPES:
                return d
        return entries_sorted[0][0]  # fallback to absolute latest

    apps = []
    for target, entries in sorted(by_target_all.items()):
        scan_dir  = _best_scan_dir(entries)
        scan_meta = parsers.load_scan(scan_dir, EPYON_ROOT)
        findings  = parsers.parse_scan_findings(scan_dir)
        apps.append({
            "name":           target,
            "monitored":      target in monitored,
            "scan_type":      scan_meta.get("scan_type", "full"),
            "critical":       scan_meta.get("critical", 0),
            "high":           scan_meta.get("high", 0),
            "medium":         scan_meta.get("medium", 0),
            "low":            scan_meta.get("low", 0),
            "tools_analyzed": scan_meta.get("tools_analyzed", []),
            "critical_sample": [
                {"tool": f.get("tool"), "id": f.get("id"), "package": f.get("package"), "title": f.get("title")}
                for f in findings.get("critical_findings", [])[:5]
            ],
        })

    # Pull programme-level metrics (uses cache if warm)
    _metrics = _build_summary_metrics()

    try:
        summary = await openai_summary.generate_global_summary(apps, metrics=_metrics)
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))
    except Exception as exc:
        raise HTTPException(502, f"OpenAI request failed: {exc}")

    return {"summary": summary, "application_count": len(apps)}


def _build_summary_metrics() -> dict:
    """Return a trimmed metrics dict for injection into AI summary prompts.
    Uses the in-memory cache when warm to avoid redundant computation."""
    import time as _time
    if _metrics_cache and (_time.time() - _metrics_cache_ts) < _METRICS_TTL:
        m = _metrics_cache
    else:
        # Cache is cold — compute now with a throwaway Response object so that
        # _sec_headers doesn't receive None.
        m = get_metrics(Response())
    if not m:
        return {}

    # Latest risk score per app
    risk_by_app: dict[str, int] = {}
    for pt in (m.get("risk_trend") or []):
        risk_by_app[pt["target"]] = pt["score"]

    # Latest secret counts per app
    secret_by_app: dict[str, dict] = {}
    for target, pts in (m.get("secret_trend_by_target") or {}).items():
        if pts:
            latest = pts[-1]
            secret_by_app[target] = {
                "verified":   latest.get("verified", 0),
                "unverified": latest.get("unverified", 0),
            }

    # Suppression counts per app
    sup_by_app: dict[str, int] = {
        t: len(v) for t, v in (m.get("suppression", {}).get("by_target") or {}).items()
    }

    sla = m.get("sla_compliance") or {}

    # Remap by_severity: rename 'breached' -> 'exceeded_sla' for cleaner AI output
    raw_by_sev = sla.get("by_severity", {})
    by_sev_for_ai = {}
    for sev, d in raw_by_sev.items():
        by_sev_for_ai[sev] = {
            "sla_days":     d.get("sla_days"),
            "within_sla":   d.get("within", 0),
            "exceeded_sla": d.get("breached", 0),
            "compliance_pct": d.get("pct"),
        }

    return {
        "sla_compliance": {
            "note":           "SLA counts are per resolved finding instance (unique tool + CVE ID + package combination). The same CVE affecting multiple packages counts once per package. 'exceeded_sla' = instances fixed AFTER the severity deadline (Critical ≤7d, High ≤30d, Medium ≤90d, Low ≤180d). Source: Epyon scan history.",
            "overall_pct":    sla.get("overall_pct"),
            "total_within_sla":   sla.get("total_within", 0),
            "total_exceeded_sla": sla.get("total_breached", 0),
            "by_severity":    by_sev_for_ai,
        },
        "recurrence_rate_pct":  m.get("recurrence_rate_pct"),
        "first_time_fix_pct":   m.get("first_time_fix_pct"),
        "total_resolved":       m.get("total_resolved", 0),
        "total_suppressed":     m.get("suppression", {}).get("total", 0),
        "risk_score_by_app":    risk_by_app,
        "secrets_by_app":       secret_by_app,
        "suppressed_by_app":    sup_by_app,
        "mttr_days":            m.get("mttr_days"),
        "mttd_days":            m.get("mttd_days"),
    }


@app.post("/api/technical-summary")
async def global_technical_summary(response: Response):
    _sec_headers(response)
    hidden    = _load_hidden_apps()
    monitored = _load_monitored_apps()
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)

    by_target_all: dict[str, list[tuple]] = {}
    for d in scan_dirs:
        meta = parsers.parse_dir_name(d.name)
        target = meta["target"]
        if target in hidden:
            continue
        by_target_all.setdefault(target, []).append((d, meta["timestamp"]))

    if not by_target_all:
        raise HTTPException(404, "No scans found")

    def _best_scan_dir(entries: list[tuple]) -> object:
        entries_sorted = sorted(entries, key=lambda x: x[1], reverse=True)
        for d, _ in entries_sorted:
            st = (parsers._read_json(d / "scan-metadata.json") or {}).get("scan_type", "")
            if st in _COMPREHENSIVE_SCAN_TYPES:
                return d
        return entries_sorted[0][0]

    apps = []
    for target, entries in sorted(by_target_all.items()):
        scan_dir  = _best_scan_dir(entries)
        scan_meta = parsers.load_scan(scan_dir, EPYON_ROOT)
        findings  = parsers.parse_scan_findings(scan_dir)
        apps.append({
            "name":           target,
            "monitored":      target in monitored,
            "scan_type":      scan_meta.get("scan_type", "full"),
            "critical":       scan_meta.get("critical", 0),
            "high":           scan_meta.get("high", 0),
            "medium":         scan_meta.get("medium", 0),
            "low":            scan_meta.get("low", 0),
            "tools_analyzed": scan_meta.get("tools_analyzed", []),
            "critical_sample": [
                {
                    "tool": f.get("tool"), "id": f.get("id"),
                    "package": f.get("package"), "version": f.get("version"),
                    "fixed_version": f.get("fixed_version"), "cvss": f.get("cvss"),
                    "title": f.get("title"),
                }
                for f in findings.get("critical_findings", [])[:15]
            ],
            "high_sample": [
                {
                    "tool": f.get("tool"), "id": f.get("id"),
                    "package": f.get("package"), "version": f.get("version"),
                    "fixed_version": f.get("fixed_version"), "cvss": f.get("cvss"),
                    "title": f.get("title"),
                }
                for f in findings.get("high_findings", [])[:10]
            ],
        })

    _metrics = _build_summary_metrics()

    try:
        summary = await openai_summary.generate_global_technical_summary(apps, metrics=_metrics)
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))
    except Exception as exc:
        raise HTTPException(502, f"OpenAI request failed: {exc}")

    return {"summary": summary, "application_count": len(apps)}


@app.post("/api/isso-summary")
async def global_isso_summary(response: Response):
    """ISSO compliance brief covering NIST/STIG controls, evidence, POA&M risk statements."""
    _sec_headers(response)
    hidden    = _load_hidden_apps()
    monitored = _load_monitored_apps()
    scan_dirs = parsers.find_scan_dirs(EPYON_ROOT)

    by_target_all: dict[str, list[tuple]] = {}
    for d in scan_dirs:
        meta   = parsers.parse_dir_name(d.name)
        target = meta["target"]
        if target in hidden:
            continue
        by_target_all.setdefault(target, []).append((d, meta["timestamp"]))

    if not by_target_all:
        raise HTTPException(404, "No scans found")

    def _best_scan_dir(entries: list[tuple]) -> object:
        entries_sorted = sorted(entries, key=lambda x: x[1], reverse=True)
        for d, _ in entries_sorted:
            st = (parsers._read_json(d / "scan-metadata.json") or {}).get("scan_type", "")
            if st in _COMPREHENSIVE_SCAN_TYPES:
                return d
        return entries_sorted[0][0]

    apps = []
    for target, entries in sorted(by_target_all.items()):
        scan_dir  = _best_scan_dir(entries)
        scan_meta = parsers.load_scan(scan_dir, EPYON_ROOT)
        findings  = parsers.parse_scan_findings(scan_dir)

        # All scan dirs for this target, newest first
        all_dirs_sorted = [d for d, _ in sorted(entries, key=lambda x: x[1], reverse=True)]

        # STIG summary: use the most recent scan that has stig-results-*.json
        stig_open = stig_pass = stig_na = stig_total = 0
        for candidate in all_dirs_sorted:
            stig_files = sorted(candidate.glob("stig-results-*.json"))
            if not stig_files:
                continue
            for sf in stig_files:
                try:
                    raw = json.loads(sf.read_text(encoding="utf-8"))
                    res = raw.get("assessments", raw) if "assessments" in raw else raw
                    stig_open  += sum(1 for v in res.values() if v.get("status") == "Open")
                    stig_pass  += sum(1 for v in res.values() if v.get("status") == "Not a Finding")
                    stig_na    += sum(1 for v in res.values() if v.get("status") in ("Not Applicable", "Not Reviewed"))
                    stig_total += len(res)
                except Exception:
                    pass
            if stig_total:
                break  # found the most recent scan with STIG data
        stig_summary = {"open": stig_open, "pass": stig_pass, "na": stig_na, "total": stig_total} if stig_total else None

        # Suppressions: use the most recent scan that has non-empty suppressed-findings.md
        suppressions: list[dict] = []
        for candidate in all_dirs_sorted:
            batch = parsers.parse_suppressed_findings(candidate)
            if batch:
                suppressions = batch
                break
        sup_sample = [
            {"type": s.get("type", ""), "severity": s.get("severity", ""), "reason": s.get("reason", "")}
            for s in suppressions[:20]
        ]

        apps.append({
            "name":           target,
            "monitored":      target in monitored,
            "scan_type":      scan_meta.get("scan_type", "full"),
            "critical":       scan_meta.get("critical", 0),
            "high":           scan_meta.get("high", 0),
            "medium":         scan_meta.get("medium", 0),
            "low":            scan_meta.get("low", 0),
            "tools_analyzed": scan_meta.get("tools_analyzed", []),
            "stig_summary":   stig_summary,
            "suppressed_count": len(suppressions),
            "suppressed_sample": sup_sample,
            "critical_sample": [
                {
                    "tool": f.get("tool"), "id": f.get("id"),
                    "package": f.get("package"), "version": f.get("version"),
                    "title": f.get("title"), "cvss": f.get("cvss"),
                }
                for f in findings.get("critical_findings", [])[:12]
            ],
            "high_sample": [
                {
                    "tool": f.get("tool"), "id": f.get("id"),
                    "package": f.get("package"), "version": f.get("version"),
                    "title": f.get("title"), "cvss": f.get("cvss"),
                }
                for f in findings.get("high_findings", [])[:8]
            ],
        })

    _metrics = _build_summary_metrics()

    try:
        summary = await openai_summary.generate_global_isso_summary(apps, metrics=_metrics)
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))
    except Exception as exc:
        raise HTTPException(502, f"OpenAI request failed: {exc}")

    return {"summary": summary, "application_count": len(apps)}


# ── Word (docx) export ─────────────────────────────────────────────────────
class _SummaryExportBody(BaseModel):
    exec_summary:  Optional[str] = None
    tech_summary:  Optional[str] = None
    isso_summary:  Optional[str] = None


@app.post("/api/export/summary-docx")
def export_summary_docx(body: _SummaryExportBody, response: Response):
    """Convert AI-generated Markdown summaries to a .docx download."""
    _sec_headers(response)
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "python-docx is not installed. Run: pip install python-docx")

    import io, re
    from datetime import datetime, timezone

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover header ──
    title_para = doc.add_heading("Epyon Security Analysis", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.runs[0]
    run.font.size  = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mr = meta.add_run(
        f"Generated: {datetime.now(timezone.utc).strftime('%B %d, %Y')}  |  AI-Assisted Report  |  Confidential"
    )
    mr.font.size  = Pt(9)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    def _add_markdown(md_text: str, section_label: str) -> None:
        """Render a Markdown string into the Word document."""
        heading_para = doc.add_heading(section_label, level=1)
        heading_para.runs[0].font.size = Pt(14)

        for line in md_text.splitlines():
            stripped = line.strip()
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=2)
            elif stripped.startswith("> "):
                # Blockquote — indent paragraph
                p = doc.add_paragraph(stripped[2:])
                p.paragraph_format.left_indent = Inches(0.4)
                p.runs[0].font.italic = True
            elif re.match(r'^[-*] ', stripped):
                _inline_paragraph(doc.add_paragraph(stripped[2:], style='List Bullet'))
            elif stripped == "":
                doc.add_paragraph()
            else:
                _inline_paragraph(doc.add_paragraph(stripped))

    def _inline_paragraph(para) -> None:
        """Apply bold/italic/code inline formatting to an existing paragraph."""
        if not para.runs:
            return
        raw = para.runs[0].text
        # python-docx Paragraph has no public clear(); remove existing runs.
        for run in list(para.runs):
            run_el = run._element
            run_el.getparent().remove(run_el)
        # Tokenise **bold**, *italic*, `code`
        tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', raw)
        for tok in tokens:
            if tok.startswith('**') and tok.endswith('**'):
                r = para.add_run(tok[2:-2])
                r.bold = True
            elif tok.startswith('*') and tok.endswith('*'):
                r = para.add_run(tok[1:-1])
                r.italic = True
            elif tok.startswith('`') and tok.endswith('`'):
                r = para.add_run(tok[1:-1])
                r.font.name = 'Courier New'
                r.font.size = Pt(9)
            else:
                para.add_run(tok)

    if body.exec_summary:
        _add_markdown(body.exec_summary, "📋 Executive Summary")
        doc.add_page_break()

    if body.tech_summary:
        _add_markdown(body.tech_summary, "🔧 Technical Summary")
        if body.isso_summary:
            doc.add_page_break()

    if body.isso_summary:
        _add_markdown(body.isso_summary, "🔒 ISSO Compliance Brief")

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    fr = footer_p.add_run(
        "Generated by Epyon Security Scanner  ·  AI-assisted — verify all findings before acting"
    )
    fr.font.size  = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fr.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    date_slug = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    filename  = f"epyon-security-report-{date_slug}.docx"
    from fastapi.responses import Response as _Resp
    return _Resp(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Per-finding fix suggestion ─────────────────────────────────

class _FindingFixRequest(BaseModel):
    id:            Optional[str] = None
    title:         Optional[str] = None
    description:   Optional[str] = None
    tool:          Optional[str] = None
    severity:      Optional[str] = None
    package:       Optional[str] = None
    version:       Optional[str] = None
    fixed_version: Optional[str] = None
    target:        Optional[str] = None
    references:    list[str]     = []


@app.post("/api/findings/fix")
async def finding_fix(body: _FindingFixRequest, response: Response):
    _sec_headers(response)
    finding = {k: v for k, v in body.model_dump().items() if v}
    if not finding:
        raise HTTPException(400, "No finding data provided")
    try:
        fix = await openai_summary.generate_fix_suggestion(finding)
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))
    except Exception as exc:
        raise HTTPException(502, f"OpenAI request failed: {exc}")
    return {"fix": fix}


# ── Jira integration ──────────────────────────────────────────

_JIRA_URL_RE    = re.compile(r"^https://[a-zA-Z0-9.\-]+/")
_JIRA_TOKEN_RE  = re.compile(r"^[A-Za-z0-9_\-]{10,}$")


@app.get("/api/jira/config")
def jira_config_get(response: Response):
    _sec_headers(response)
    cfg   = jira_client.read_config()
    token = cfg.get("api_token", "")
    masked = re.sub(r"(?<=.{4}).(?=.{4})", "*", token) if len(token) > 8 else ("*" * len(token))
    return {
        "token_set":       bool(token),
        "token_hint":      masked,
        "base_url":        cfg.get("base_url", ""),
        "email":           cfg.get("email", ""),
        "project_key":     cfg.get("project_key", ""),
        "issue_type":      cfg.get("issue_type", "Bug"),
        "done_transition": cfg.get("done_transition", "Done"),
        "min_severity":    cfg.get("min_severity", "high"),
        "auto_close":      cfg.get("auto_close", False),
        "create_on_new":   cfg.get("create_on_new", False),
        "_from_env":       cfg.get("_from_env", False),
    }


@app.post("/api/jira/config")
async def jira_config_post(request: Request, response: Response):
    _sec_headers(response)
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "Invalid JSON")

    cfg = jira_client.read_config()

    base_url = (body.get("base_url") or "").strip()
    if base_url:
        if not base_url.startswith("https://"):
            raise HTTPException(400, "base_url must use HTTPS")
        if re.search(r"[;&|`$\(\)\n\r<>'\"\\]", base_url):
            raise HTTPException(400, "base_url contains invalid characters")
        cfg["base_url"] = base_url

    email = (body.get("email") or "").strip()
    if email:
        if not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email):
            raise HTTPException(400, "Invalid email address")
        cfg["email"] = email

    new_token = (body.get("api_token") or "").strip()
    if new_token and new_token != "KEEP_EXISTING":
        cfg["api_token"] = new_token

    if body.get("project_key") is not None:
        pk = (body["project_key"] or "").strip().upper()
        if pk and not re.match(r"^[A-Z][A-Z0-9]{0,9}$", pk):
            raise HTTPException(400, "project_key must be 1-10 uppercase alphanumeric characters")
        cfg["project_key"] = pk

    for str_field in ("issue_type", "done_transition"):
        if body.get(str_field) is not None:
            val = (body[str_field] or "").strip()
            if val and re.search(r"[<>\"';&|`$]", val):
                raise HTTPException(400, f"{str_field} contains invalid characters")
            cfg[str_field] = val

    if body.get("min_severity") in ("critical", "high", "medium", "low"):
        cfg["min_severity"] = body["min_severity"]

    if "auto_close" in body:
        cfg["auto_close"] = bool(body["auto_close"])
    if "create_on_new" in body:
        cfg["create_on_new"] = bool(body["create_on_new"])

    jira_client.write_config(cfg)
    _audit(request, "jira_config_changed")
    return {"ok": True}


@app.post("/api/jira/test")
async def jira_test(response: Response):
    _sec_headers(response)
    cfg = jira_client.read_config()
    result = await jira_client.test_connection(cfg)
    return result


@app.get("/api/jira/tickets")
def jira_tickets_list(response: Response):
    _sec_headers(response)
    tmap = jira_client.read_ticket_map()
    # Return sorted list, newest created first
    entries = list(tmap.values())
    entries.sort(key=lambda e: e.get("created_at", ""), reverse=True)
    return {"total": len(entries), "tickets": entries}


@app.post("/api/jira/sync/{app_name}")
async def jira_sync_app(app_name: str, response: Response):
    """Manually trigger Jira reconciliation for a specific application."""
    _sec_headers(response)
    if not _SAFE_ID_RE.match(app_name):
        raise HTTPException(400, "Invalid app_name")

    cfg = jira_client.read_config()
    if not cfg.get("api_token"):
        raise HTTPException(400, "Jira is not configured — set api_token first")

    all_dirs = parsers.find_scan_dirs(EPYON_ROOT)
    target_dirs = sorted(
        [d for d in all_dirs if parsers.parse_dir_name(d.name)["target"] == app_name],
        key=lambda d: d.name,
        reverse=True,
    )
    if len(target_dirs) < 2:
        raise HTTPException(400, "Need at least two scans for this application to compare findings")

    current_raw  = (parsers.load_enriched_findings(target_dirs[0])
                    or parsers.parse_scan_findings(target_dirs[0]))
    previous_raw = (parsers.load_enriched_findings(target_dirs[1])
                    or parsers.parse_scan_findings(target_dirs[1]))

    ticket_map = jira_client.read_ticket_map()
    result = await jira_client.reconcile_and_save(
        app_name,
        jira_client.flatten_findings(current_raw),
        jira_client.flatten_findings(previous_raw),
        cfg,
    )
    return result


# ── SPA / static file serving ─────────────────────────────────
# Serves web/static/ for JS, CSS, and other assets.

if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

    @app.get("/{full_path:path}", include_in_schema=False)
    def spa_fallback(full_path: str):
        candidate = STATIC_DIR / full_path
        if candidate.is_file():
            return FileResponse(str(candidate))
        index = STATIC_DIR / "index.html"
        if index.exists():
            return FileResponse(str(index))
        return JSONResponse({"detail": "Frontend not found"}, status_code=503)
else:
    @app.get("/{full_path:path}", include_in_schema=False)
    def spa_not_found(full_path: str):
        return JSONResponse(
            {"detail": "Static directory not found"},
            status_code=503,
        )
